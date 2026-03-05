#!/usr/bin/env python
# -*- coding: utf-8 -*-

"""
文档处理器 - 整合格式调整和智能应答功能
"""

import sys
import os
import re
import traceback
import comtypes.client
import logging
from datetime import datetime
from pathlib import Path

# PyQt5 imports
from PyQt5.QtCore import QThread, pyqtSignal

# Document processing imports
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph
import win32com.client
from clause_utils import extract_numbered_paragraphs, get_prefix_and_level, analyze_hierarchy, find_minimal_clauses
from similarity import calculate_similarity, find_best_match, SIMILARITY_THRESHOLD


# ============== 回应条款模板管理（精准应答版）==============

def get_template_folder_path():
    """
    获取回应条款文件夹路径
    - 打包后：exe同级目录下的"回应条款"文件夹
    - 开发环境：项目根目录下的"回应条款"文件夹
    """
    if getattr(sys, 'frozen', False):
        app_dir = os.path.dirname(sys.executable)
    else:
        app_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    return os.path.join(app_dir, '回应条款')

def get_product_list_from_folder():
    """
    动态扫描“回应条款”文件夹，生成产品列表。
    - 自动忽略“其他.docx”和Word临时文件（~$开头）。
    - 返回按字母排序的产品名称列表。
    """
    template_dir = get_template_folder_path()
    if not os.path.isdir(template_dir):
        return []
    
    products = []
    try:
        for filename in os.listdir(template_dir):
            if filename.endswith('.docx') and not filename.startswith('~$'):
                product_name = os.path.splitext(filename)[0]
                if product_name != '其他':
                    products.append(product_name)
    except FileNotFoundError:
        # 如果文件夹不存在，返回空列表
        return []
    
    products.sort()
    return products

# 产品列表（从“回应条款”文件夹动态生成）
PRODUCT_LIST = get_product_list_from_folder()

# 全局模板存储
# 新结构: {产品名: [{'clause_text': 条款内容, 'response_text': 应答内容,
#                   'template_file': 模板文件路径, 'element_indices': [(type, idx), ...]}, ...]}
_product_templates = {}
_default_template = []   # 其他.docx 解析后的条款-应答对
_templates_loaded = False
_default_response = "我方完全满足招标文件要求。"


def parse_template_file(file_path):
    """
    解析模板文件，提取条款-应答对（包含元素索引信息，用于复制表格和图片）

    支持的标记格式：
    - [条款N] / [应答N]
    - [条款N] / [应答内容N]
    - 条款N： / 应答内容N：

    返回：[{'clause_text': 条款内容, 'response_text': 应答内容,
            'template_file': 模板文件路径, 'element_indices': [(type, idx), ...]}, ...]
    """
    try:
        doc = Document(file_path)

        # 获取文档body中所有元素（段落和表格）
        body_elements = []
        if hasattr(doc, '_body') and hasattr(doc._body, '_body'):
            for idx, element in enumerate(doc._body._body):
                tag = element.tag.split('}')[-1] if '}' in element.tag else element.tag
                # 检查元素是否包含图片
                elem_xml = element.xml if hasattr(element, 'xml') else ''
                has_image = 'r:embed' in elem_xml
                if tag == 'p':
                    # 获取段落文本
                    para_idx = sum(1 for e in list(doc._body._body)[:idx] if (e.tag.split('}')[-1] if '}' in e.tag else e.tag) == 'p')
                    if para_idx < len(doc.paragraphs):
                        text = doc.paragraphs[para_idx].text.strip()
                        body_elements.append({'type': 'paragraph', 'index': idx, 'text': text, 'para_idx': para_idx, 'has_image': has_image})
                elif tag == 'tbl':
                    body_elements.append({'type': 'table', 'index': idx, 'text': '[TABLE]', 'has_image': has_image})

        clause_pairs = []
        current_clause = []
        current_response = []
        current_response_elements = []  # 保存应答部分的元素索引
        current_is_superior = False  # 记录当前应答是否有[优于]标记
        in_clause = False
        in_response = False

        # 支持多种条款标记格式：[条款N] 或 条款N：
        clause_pattern = re.compile(r'^(\[条款\d+\]|条款\d+[：:])$')
        # 支持多种应答标记格式：[应答N] 或 [应答内容N] 或 应答内容N：
        response_pattern = re.compile(r'^(\[应答\d+\]|\[应答内容\d+\]|应答内容?\d+[：:])$')

        for elem in body_elements:
            text = elem['text']
            has_image = elem.get('has_image', False)

            if clause_pattern.match(text):
                # 保存上一个条款-应答对
                if current_clause and current_response:
                    clause_pairs.append({
                        'clause_text': '\n'.join(current_clause),
                        'response_text': '\n'.join(current_response),
                        'template_file': file_path,
                        'element_indices': current_response_elements.copy(),
                        'is_superior': current_is_superior
                    })
                current_clause = []
                current_response = []
                current_response_elements = []
                current_is_superior = False
                in_clause = True
                in_response = False
            elif response_pattern.match(text):
                in_clause = False
                in_response = True
            elif in_clause:
                if text:  # 只添加非空文本
                    current_clause.append(text)
            elif in_response:
                # 检测[优于]标记
                if text == '[优于]':
                    current_is_superior = True
                # 跳过 [优于] [满足] 等标记（但保留其他内容）
                if text and not re.match(r'^\[.+\]$', text):
                    current_response.append(text if elem['type'] == 'paragraph' else '')
                # 保存元素索引（包括表格和包含图片的元素）
                should_save = (
                    elem['type'] == 'table' or
                    has_image or  # 包含图片的元素必须保存
                    (text and not re.match(r'^\[.+\]$', text))
                )
                if should_save:
                    current_response_elements.append((elem['type'], elem['index']))

        # 保存最后一个条款-应答对
        if current_clause and current_response:
            clause_pairs.append({
                'clause_text': '\n'.join(current_clause),
                'response_text': '\n'.join(current_response),
                'template_file': file_path,
                'element_indices': current_response_elements.copy(),
                'is_superior': current_is_superior
            })

        return clause_pairs
    except Exception as e:
        print(f"解析模板文件失败 {file_path}: {e}")
        import traceback
        traceback.print_exc()
        return []


def load_all_templates():
    """
    加载所有产品模板，解析为条款-应答对
    """
    global _product_templates, _default_template, _templates_loaded, _default_response

    if _templates_loaded:
        return _product_templates

    template_dir = get_template_folder_path()
    _product_templates = {}
    _default_template = []

    if not os.path.exists(template_dir):
        print(f"回应条款文件夹不存在: {template_dir}")
        _templates_loaded = True
        return _product_templates

    # 递归扫描所有.docx文件
    for root, dirs, files in os.walk(template_dir):
        for filename in files:
            if filename.endswith('.docx') and not filename.startswith('~$'):
                file_path = os.path.join(root, filename)
                product_name = os.path.splitext(filename)[0]

                # 解析模板文件
                clause_pairs = parse_template_file(file_path)

                if clause_pairs:
                    if product_name == '其他':
                        _default_template = clause_pairs
                        print(f"已加载默认模板: 其他 ({len(clause_pairs)}个条款)")
                    else:
                        _product_templates[product_name] = clause_pairs
                        print(f"已加载产品模板: {product_name} ({len(clause_pairs)}个条款)")

    _templates_loaded = True
    print(f"共加载 {len(_product_templates)} 个产品模板")
    return _product_templates


def get_available_products():
    """获取可用的产品列表"""
    load_all_templates()
    return [p for p in PRODUCT_LIST if p in _product_templates]


def match_clause_with_product(clause_text, product_name):
    """
    使用指定产品的模板匹配条款

    匹配逻辑：
    1. 先用指定产品的模板进行相似度匹配
    2. 如果相似度不够，再用"其他.docx"模板匹配
    3. 如果还是没有匹配，返回默认应答

    参数：
        clause_text - 招标文档中的条款内容
        product_name - 用户选择的产品名称
    返回：
        (matched, response_data, similarity, source)
        matched: 是否匹配成功
        response_data: 应答数据，可能是：
            - 字典格式: {'response_text': 应答文本, 'template_file': 模板路径, 'element_indices': 元素索引}
            - 字符串格式: 默认应答文本
        similarity: 相似度
        source: 来源（产品名/其他/默认）
    """
    global _product_templates, _default_template, _default_response

    load_all_templates()

    best_similarity = 0.0

    # 1. 先用指定产品的模板匹配
    if product_name and product_name in _product_templates:
        matched, response_data, similarity, _ = find_best_match(
            clause_text, _product_templates[product_name]
        )
        best_similarity = max(best_similarity, similarity)
        if matched:
            return True, response_data, similarity, product_name

    # 2. 如果相似度不够，用"其他"模板匹配
    if _default_template:
        matched, response_data, similarity, _ = find_best_match(
            clause_text, _default_template
        )
        best_similarity = max(best_similarity, similarity)
        if matched:
            return True, response_data, similarity, "其他"

    # 3. 如果还是没有匹配，返回默认应答（返回最高相似度供调试）
    return False, _default_response, best_similarity, "默认"


def reload_templates():
    """重新加载模板"""
    global _templates_loaded
    _templates_loaded = False
    return load_all_templates()


# 兼容旧接口
def load_response_templates():
    """兼容旧接口"""
    return load_all_templates()


def match_template_for_clause(clause_text, product_name=None):
    """
    兼容旧接口，支持新的产品匹配
    """
    if product_name:
        matched, response_data, similarity, source = match_clause_with_product(clause_text, product_name)
        # 从新格式中提取应��文本
        if isinstance(response_data, dict):
            response_text = response_data.get('response_text', _default_response)
        else:
            response_text = response_data
        return source if matched else None, [response_text]
    else:
        # 旧逻辑：关键词匹配
        load_all_templates()
        for product_name, pairs in _product_templates.items():
            if product_name in clause_text:
                # 返回第一个条款的应答作为默认
                if pairs:
                    item = pairs[0]
                    if isinstance(item, dict):
                        return product_name, [item.get('response_text', _default_response)]
                    else:
                        return product_name, [item[1]]
        return None, [_default_response]


# ============== 文档处理器 ==============

class DocumentProcessor(QThread):
    """文档处理线程 - 完全按照原始逻辑实现"""
    progress = pyqtSignal(int)
    log = pyqtSignal(str)
    finished = pyqtSignal(str)
    error = pyqtSignal(str)

    def __init__(self, input_file, process_mode, format_options=None, selected_product=None):
        super().__init__()
        self.input_file = input_file
        self.process_mode = process_mode  # "format", "response", "both"
        self.output_file = None
        self.selected_product = selected_product  # 用户选择的产品（用于精准应答）
        self.format_options = format_options or {
            'outline': True,
            'numbering': True,
            'image': True,
            'table': True,
            'keyword': True,
            'symbol': True,
            'header_footer': True
        }
        # 初始化详细日志
        self._init_file_logger()

    def _init_file_logger(self):
        """初始化文件日志记录器"""
        # 确保logs目录存在
        log_dir = os.path.join(os.path.dirname(os.path.dirname(__file__)), 'logs')
        os.makedirs(log_dir, exist_ok=True)

        # 创建日志文件名（包含时间戳）
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        input_name = os.path.splitext(os.path.basename(self.input_file))[0]
        log_file = os.path.join(log_dir, f'{timestamp}_{input_name}.log')

        # 配置logger
        self.file_logger = logging.getLogger(f'DocProcessor_{timestamp}')
        self.file_logger.setLevel(logging.DEBUG)
        self.file_logger.handlers.clear()

        handler = logging.FileHandler(log_file, encoding='utf-8')
        handler.setFormatter(logging.Formatter('%(asctime)s [%(levelname)s] %(message)s'))
        self.file_logger.addHandler(handler)

        self.file_logger.info(f'========== 开始处理文档 ==========')
        self.file_logger.info(f'输入文件: {self.input_file}')
        self.file_logger.info(f'处理模式: {self.process_mode}')
        self.file_logger.info(f'选择产品: {self.selected_product or "未选择（常规应答）"}')

    def _count_images(self, doc, location=''):
        """统计文档中的图片数量"""
        drawing = 0
        pict = 0
        for i, p in enumerate(doc.paragraphs):
            xml = p._p.xml
            if '<w:drawing' in xml:
                drawing += 1
                self.file_logger.debug(f'  段落{i}: 包含drawing图片')
            if '<w:pict' in xml:
                pict += 1
                self.file_logger.debug(f'  段落{i}: 包含pict图片')
        self.file_logger.info(f'[图片统计] {location}: drawing={drawing}, pict={pict}, 总计={drawing+pict}')
        return drawing, pict
        
    def run(self):
        try:
            if self.process_mode == "format":
                self.process_format_only()
            elif self.process_mode == "response":
                self.process_response_only()
            elif self.process_mode == "both":
                self.process_both()
        except Exception as e:
            self.error.emit(f"处理出错：{str(e)}\n{traceback.format_exc()}")
    
    def process_format_only(self):
        """仅进行格式调整"""
        self.log.emit("开始文档格式调整...")
        self.progress.emit(5)
        
        # 第一步：doc转docx（如果需要）
        docx_file = self.convert_doc_to_docx()
        self.progress.emit(15)
        
        # 第二步：打开文档
        doc = Document(docx_file)
        self.log.emit(f"成功打开文档：{os.path.basename(docx_file)}")
        self.progress.emit(18)
        
        # 执行格式调整的完整流程
        doc = self.execute_format_adjustment(doc)
        
        # 保存文档
        base_name = os.path.splitext(os.path.basename(self.input_file))[0]
        self.output_file = os.path.join(os.path.dirname(self.input_file), f"{base_name}-格式调整.docx")
        doc.save(self.output_file)
        self.log.emit(f"文档保存成功：{os.path.basename(self.output_file)}")
        self.log.emit(f"保存路径：{os.path.abspath(self.output_file)}")
        self.progress.emit(100)
        
        self.finished.emit(self.output_file)
    
    def process_response_only(self):
        """仅进行智能应答"""
        self.log.emit("开始智能应答处理...")
        self.progress.emit(5)
        
        # 打开文档
        doc = Document(self.input_file)
        self.log.emit(f"成功打开文档：{os.path.basename(self.input_file)}")
        self.progress.emit(10)
        
        # 执行应答处理的完整流程
        self.execute_response_processing(doc)
        
        # 保存文档
        base_name = os.path.splitext(os.path.basename(self.input_file))[0]
        self.output_file = os.path.join(os.path.dirname(self.input_file), f"{base_name}-智能应答.docx")
        doc.save(self.output_file)
        self.log.emit(f"文档保存成功：{os.path.basename(self.output_file)}")
        self.log.emit(f"保存路径：{os.path.abspath(self.output_file)}")
        self.progress.emit(100)
        
        self.finished.emit(self.output_file)
    
    def process_both(self):
        """同时进行格式调整和智能应答"""
        self.log.emit("开始格式调整+智能应答处理...")
        self.file_logger.info('========== process_both 开始 ==========')
        self.progress.emit(5)

        # 第一步：doc转docx（如果需要）
        docx_file = self.convert_doc_to_docx()
        self.progress.emit(10)

        # 第二步：打开文档
        doc = Document(docx_file)
        self.log.emit(f"成功打开文档：{os.path.basename(docx_file)}")
        self.file_logger.info(f'[STEP] 打开文档: {docx_file}')
        self._count_images(doc, '打开文档后')
        self.progress.emit(15)

        # 第三步：先执行格式调整
        self.log.emit("第一阶段：执行格式调整...")
        self.file_logger.info('========== 第一阶段：格式调整 ==========')
        doc = self.execute_format_adjustment(doc, progress_offset=15, progress_range=40)
        self._count_images(doc, '格式调整完成后')

        # 第四步：再执行应答处理
        self.log.emit("第二阶段：执行智能应答...")
        self.file_logger.info('========== 第二阶段：智能应答 ==========')
        self.execute_response_processing(doc, progress_offset=55, progress_range=40)

        # 统计输出文档图片数量
        self._count_images(doc, '智能应答完成后')
        out_drawing = sum(1 for p in doc.paragraphs if '<w:drawing' in p._p.xml)
        out_pict = sum(1 for p in doc.paragraphs if '<w:pict' in p._p.xml)
        self.log.emit(f"输出文档图片统计: drawing={out_drawing}, pict={out_pict}")

        # 保存文档
        base_name = os.path.splitext(os.path.basename(self.input_file))[0]
        self.output_file = os.path.join(os.path.dirname(self.input_file), f"{base_name}-格式调整+智能应答.docx")
        doc.save(self.output_file)
        self.log.emit(f"文档保存成功：{os.path.basename(self.output_file)}")
        self.log.emit(f"保存路径：{os.path.abspath(self.output_file)}")
        self.progress.emit(100)
        
        self.finished.emit(self.output_file)
    
    def convert_doc_to_docx(self):
        """将.doc文件转换为.docx文件 - 完全按照V1.0逻辑"""
        self.file_logger.info('[convert_doc_to_docx] 开始')
        if self.input_file.lower().endswith('.docx'):
            self.log.emit("文件已是.docx格式，无需转换")
            self.file_logger.info('[convert_doc_to_docx] 文件已是.docx格式')
            return self.input_file

        self.log.emit("正在转换.doc文件为.docx格式...")
        self.file_logger.info(f'[convert_doc_to_docx] 转换.doc到.docx: {self.input_file}')

        try:
            # 使用win32com转换
            word = win32com.client.Dispatch("Word.Application")
            word.Visible = False

            doc = word.Documents.Open(os.path.abspath(self.input_file))

            # 生成输出文件名
            base_name = os.path.splitext(self.input_file)[0]
            docx_file = f"{base_name}.docx"

            # 保存为docx格式 (16 = wdFormatXMLDocument)
            doc.SaveAs2(os.path.abspath(docx_file), FileFormat=16)
            doc.Close()
            word.Quit()

            self.log.emit(f"转换完成：{os.path.basename(docx_file)}")
            self.file_logger.info(f'[convert_doc_to_docx] 转换成功: {docx_file}')
            return docx_file

        except Exception as e:
            self.log.emit(f"转换失败，尝试直接处理原文件：{str(e)}")
            self.file_logger.error(f'[convert_doc_to_docx] 转换失败: {e}')
            return self.input_file
    
    def execute_format_adjustment(self, doc, progress_offset=18, progress_range=82):
        """执行完整的格式调整流程 - 支持细分功能选项"""
        # 计算需要执行的步骤数
        enabled_steps = []

        # 始终执行的基础步骤
        enabled_steps.append('basic')  # 基础格式处理
        enabled_steps.append('margins')  # 页面格式
        enabled_steps.append('styles')  # 创建样式
        enabled_steps.append('cleanup')  # 最终空格清理

        # 根据选项添加步骤
        if self.format_options.get('numbering', True):
            enabled_steps.insert(0, 'numbering')  # 自动序号转文本
        if self.format_options.get('outline', True):
            enabled_steps.append('outline')  # 大纲生成（标题样式）
            enabled_steps.append('body_styles')  # 正文样式（与大纲配合）
        if self.format_options.get('table', True):
            enabled_steps.append('table')  # 表格调整
        if self.format_options.get('image', True):
            enabled_steps.append('image')  # 图片调整
        if self.format_options.get('keyword', True) or self.format_options.get('symbol', True):
            enabled_steps.append('highlight')  # 关键词/符号标记
        if self.format_options.get('header_footer', True):
            enabled_steps.append('header_footer')  # 删除页眉页脚

        step_size = progress_range / len(enabled_steps)
        current_step = 0

        # 自动序号转文本
        if 'numbering' in enabled_steps:
            self.log.emit("转换自动编号为文本编号...")
            self.file_logger.info('[STEP] 开始: 转换自动编号为文本编号')
            doc = self.convert_auto_numbering_to_text(doc)
            self._count_images(doc, '自动编号转换后')
            current_step += 1
            self.progress.emit(int(progress_offset + step_size * current_step))

        # 基础格式处理
        if 'basic' in enabled_steps:
            self.log.emit("基础格式处理...")
            self.file_logger.info('[STEP] 开始: 基础格式处理')
            self.basic_format_cleanup(doc)
            self._count_images(doc, '基础格式处理后')
            current_step += 1
            self.progress.emit(int(progress_offset + step_size * current_step))

        # 设置页面格式
        if 'margins' in enabled_steps:
            self.log.emit("设置页面格式...")
            self.file_logger.info('[STEP] 开始: 设置页面格式')
            self.set_page_margins(doc)
            current_step += 1
            self.progress.emit(int(progress_offset + step_size * current_step))

        # 创建和设置样式
        if 'styles' in enabled_steps:
            self.log.emit("创建和设置样式...")
            self.file_logger.info('[STEP] 开始: 创建和设置样式')
            doc = self.create_custom_styles(doc)
            current_step += 1
            self.progress.emit(int(progress_offset + step_size * current_step))

        # 识别序号和应用标题样式（大纲生成）
        if 'outline' in enabled_steps:
            self.log.emit("识别序号和应用标题样式（大纲生成）...")
            self.file_logger.info('[STEP] 开始: 识别序号和应用标题样式')
            doc = self.apply_heading_styles(doc)
            self._count_images(doc, '应用标题样式后')
            current_step += 1
            self.progress.emit(int(progress_offset + step_size * current_step))

        # 应用正文样式
        if 'body_styles' in enabled_steps:
            self.log.emit("应用正文样式...")
            self.file_logger.info('[STEP] 开始: 应用正文样式')
            doc = self.apply_body_styles(doc)
            self._count_images(doc, '应用正文样式后')
            current_step += 1
            self.progress.emit(int(progress_offset + step_size * current_step))

        # 处理表格格式
        if 'table' in enabled_steps:
            self.log.emit("处理表格格式...")
            self.file_logger.info('[STEP] 开始: 处理表格格式')
            self.format_tables(doc)
            current_step += 1
            self.progress.emit(int(progress_offset + step_size * current_step))

        # 处理图片格式
        if 'image' in enabled_steps:
            self.log.emit("处理图片格式...")
            self.file_logger.info('[STEP] 开始: 处理图片格式')
            self.format_images(doc)
            current_step += 1
            self.progress.emit(int(progress_offset + step_size * current_step))

        # 最终全文空格清理
        if 'cleanup' in enabled_steps:
            self.log.emit("最终全文空格清理...")
            self.file_logger.info('[STEP] 开始: 最终全文空格清理')
            self.final_cleanup_all_spaces(doc)
            current_step += 1
            self.progress.emit(int(progress_offset + step_size * current_step))

        # 关键词标青绿色和符号标红处理
        if 'highlight' in enabled_steps:
            keyword_enabled = self.format_options.get('keyword', True)
            symbol_enabled = self.format_options.get('symbol', True)
            if keyword_enabled and symbol_enabled:
                self.log.emit("关键词标青绿色和符号标红处理...")
            elif keyword_enabled:
                self.log.emit("关键词标青绿色处理...")
            else:
                self.log.emit("符号标红处理...")
            self.highlight_keywords_with_options(doc, keyword_enabled, symbol_enabled)
            current_step += 1
            self.progress.emit(int(progress_offset + step_size * current_step))

        # 删除页眉页脚
        if 'header_footer' in enabled_steps:
            self.log.emit("删除页眉页脚...")
            self.remove_headers_and_footers(doc)
            current_step += 1
            self.progress.emit(int(progress_offset + step_size * current_step))

        return doc
    
    def convert_auto_numbering_to_text(self, doc):
        """将自动编号转换为文本编号 - 改进版本，正确处理项目符号"""
        self.file_logger.info('[convert_auto_numbering_to_text] 开始')
        self.log.emit("- 检查并转换自动编号...")

        # 首先检查是否存在自动编号
        has_auto_numbering = False
        auto_num_count = 0
        for para in doc.paragraphs:
            if self.paragraph_may_have_auto_numbering(para):
                has_auto_numbering = True
                auto_num_count += 1

        self.file_logger.info(f'[convert_auto_numbering_to_text] 检测到 {auto_num_count} 个自动编号段落')

        if not has_auto_numbering:
            self.log.emit("- 未检测到自动编号，跳过转换")
            self.file_logger.info('[convert_auto_numbering_to_text] 无自动编号，跳过')
            return doc

        self.log.emit("- 检测到自动编号，开始转换...")
        self.file_logger.info('[编号转换] 开始win32com处理')

        try:
            # 使用系统临时目录和唯一文件名
            import tempfile
            import uuid
            temp_dir = tempfile.gettempdir()
            temp_file = os.path.join(temp_dir, f"numbering_{uuid.uuid4().hex[:8]}.docx")
            self.file_logger.info(f'[编号转换] 准备保存临时文件: {temp_file}')
            doc.save(temp_file)
            self.file_logger.info(f'[编号转换] 临时文件已保存: {temp_file}')

            # 使用win32com处理自动编号
            word = win32com.client.Dispatch("Word.Application")
            word.Visible = False
            self.file_logger.info('[编号转换] Word应用已启动')

            # 用win32com打开临时文件
            word_doc = word.Documents.Open(os.path.abspath(temp_file))
            self.file_logger.info('[编号转换] 文档已打开')

            # 收集需要处理的段落信息（从后向前处理，避免索引变化）
            para_info_list = []
            skipped_image_count = 0
            for i, para in enumerate(word_doc.Paragraphs):
                try:
                    list_format = para.Range.ListFormat
                    list_type = list_format.ListType
                    if list_type != 0:
                        # 检查段落是否包含图片，如果包含则跳过
                        has_image = para.Range.InlineShapes.Count > 0
                        if has_image:
                            skipped_image_count += 1
                            self.file_logger.info(f'[编号转换] 跳过含图片段落{i}: InlineShapes={para.Range.InlineShapes.Count}')
                            continue

                        # 直接获取列表编号字符串（如 "A." "1." "(1)" 等）
                        list_string = list_format.ListString

                        para_info_list.append({
                            'index': i,
                            'list_string': list_string,
                            'list_type': list_type
                        })
                except Exception as e:
                    self.file_logger.error(f'[编号转换] 收集段落{i}信息失败: {e}')

            self.file_logger.info(f'[编号转换] 收集到 {len(para_info_list)} 个待处理段落')

            # 从后向前处理段落，在段落开头插入编号文本
            converted_count = 0
            for info in reversed(para_info_list):
                try:
                    para = word_doc.Paragraphs(info['index'] + 1)  # VBA索引从1开始
                    list_string = info['list_string']

                    if list_string:
                        # 获取段落范围
                        rng = para.Range

                        # 移除列表格式（但保持段落独立）
                        rng.ListFormat.RemoveNumbers()

                        # 在段落开头插入编号文本
                        start_rng = para.Range
                        start_rng.Collapse(1)  # wdCollapseStart = 1
                        start_rng.InsertBefore(list_string)

                        converted_count += 1
                except Exception as e:
                    self.file_logger.error(f'[编号转换] 处理段落{info["index"]}失败: {e}')

            self.log.emit(f"- 已转换 {converted_count} 个编号段落")
            if skipped_image_count > 0:
                self.file_logger.info(f'[编号转换] 跳过了 {skipped_image_count} 个含图片的段落')

            # 取消所有字段链接
            try:
                word_doc.Content.Select()
                word_doc.Application.Selection.Fields.Unlink()
            except:
                pass

            # 保存并关闭
            word_doc.Save()
            word_doc.Close()
            word.Quit()

            # 重新加载文档
            self.file_logger.info('[convert_auto_numbering_to_text] 重新加载文档')
            doc = Document(temp_file)

            # 统计重新加载后的图片
            drawing_count = sum(1 for p in doc.paragraphs if '<w:drawing' in p._p.xml)
            pict_count = sum(1 for p in doc.paragraphs if '<w:pict' in p._p.xml)
            self.file_logger.info(f'[convert_auto_numbering_to_text] 重新加载后图片: drawing={drawing_count}, pict={pict_count}')

            # 删除临时文件
            try:
                os.remove(temp_file)
            except:
                pass

            self.log.emit("- 自动编号转换完成")
            self.file_logger.info('[convert_auto_numbering_to_text] 完成')
            return doc

        except Exception as e:
            self.log.emit(f"- 自动编号转换失败，继续处理：{str(e)}")
            self.file_logger.error(f'[编号转换] 异常: {e}')
            import traceback
            self.file_logger.error(f'[编号转换] 堆栈: {traceback.format_exc()}')
            try:
                if 'temp_file' in locals() and os.path.exists(temp_file):
                    os.remove(temp_file)
            except:
                pass
            try:
                if 'word' in locals():
                    word.Quit()
            except:
                pass
            return doc
    
    def paragraph_may_have_auto_numbering(self, paragraph):
        """检查段落是否可能包含自动编号 - 完全按照V1.0逻辑"""
        if not paragraph.text.strip():
            return False
        
        # 检查段落XML是否包含列表相关元素
        xml = paragraph._p.xml
        return any(tag in xml for tag in ['<w:numPr', '<w:numId', '<w:ilvl'])
    
    def basic_format_cleanup(self, doc):
        """基础格式清理 - 完全按照V1.0逻辑- 改进版本，增加删除分节符功能"""
        self.file_logger.info('[basic_format_cleanup] 开始')
        self.log.emit("- 删除文档中的分节符...")
        self.remove_section_breaks(doc)

        self.log.emit("- 删除多余空行和空格...")

        # 删除空段落和多余空格
        paragraphs_to_remove = []
        special_kept = 0
        for i, para in enumerate(doc.paragraphs):
            if para.text.strip() == "":
                # 检查是否包含特殊元素（表格、图片等）
                if not self.has_special_elements(para):
                    paragraphs_to_remove.append(para)
                else:
                    special_kept += 1
                    self.file_logger.debug(f'[basic_format_cleanup] 保留含特殊元素的空段落{i}')
            else:
                # 清理段落中的多余空格
                self.clean_paragraph_spaces(para)

        # 删除空段落
        for para in paragraphs_to_remove:
            try:
                para._element.getparent().remove(para._element)
            except:
                pass

        self.log.emit(f"- 已删除 {len(paragraphs_to_remove)} 个空段落")
        self.file_logger.info(f'[basic_format_cleanup] 删除{len(paragraphs_to_remove)}个空段落, 保留{special_kept}个含特殊元素段落')

        try:
            # 验证分节符删除效果
            remaining_sectPrs = doc._element.xpath(".//w:pPr/w:sectPr")
            if remaining_sectPrs:
                self.log.emit(f"- 警告：仍有 {len(remaining_sectPrs)} 个分节符无法删除")
                # 可选：尝试Win32com备用方案
                # self.remove_section_breaks_win32com_fallback(self.input_file)
            else:
                self.log.emit("- 分节符删除完成")
        except Exception as e:
            self.log.emit(f"- 无法验证分节符删除效果: {e}")
    
    def remove_section_breaks(self, doc):
        """删除文档中的分节符 - 修复版本"""
        try:
            removed_count = 0
            
            # 方法1：使用文档级别的XPath（推荐）- 最高效
            try:
                sectPrs = doc._element.xpath(".//w:pPr/w:sectPr")
                for sectPr in sectPrs:
                    try:
                        parent = sectPr.getparent()
                        if parent is not None:
                            parent.remove(sectPr)
                            removed_count += 1
                    except Exception as e:
                        self.log.emit(f"删除单个分节符失败: {e}")
                        continue
                
                self.log.emit(f"- 方法1成功删除 {removed_count} 个分节符")
                
            except Exception as e1:
                self.log.emit(f"- 方法1失败，尝试方法2: {e1}")
                
                # 方法2：段落级别迭代（备用方案）
                try:
                    for paragraph in doc.paragraphs:
                        if hasattr(paragraph, '_p') and paragraph._p is not None:
                            sectPrs = paragraph._p.xpath("./w:pPr/w:sectPr")
                            for sectPr in sectPrs:
                                try:
                                    parent = sectPr.getparent()
                                    if parent is not None:
                                        parent.remove(sectPr)
                                        removed_count += 1
                                except Exception as e:
                                    continue
                    
                    self.log.emit(f"- 方法2成功删除 {removed_count} 个分节符")
                    
                except Exception as e2:
                    self.log.emit(f"- 方法2也失败，尝试方法3: {e2}")
                    
                    # 方法3：直接遍历XML元素（最后备用）
                    try:
                        from docx.oxml.ns import qn
                        
                        # 查找所有包含分节符的段落属性
                        all_pPrs = doc._element.xpath(".//w:pPr")
                        for pPr in all_pPrs:
                            sectPr_elements = pPr.findall(qn('w:sectPr'))
                            for sectPr in sectPr_elements:
                                try:
                                    pPr.remove(sectPr)
                                    removed_count += 1
                                except Exception as e:
                                    continue
                        
                        self.log.emit(f"- 方法3成功删除 {removed_count} 个分节符")
                        
                    except Exception as e3:
                        self.log.emit(f"- 所有方法都失败: {e1}, {e2}, {e3}")
                        self.log.emit("- 建议使用Win32com方法处理")
            
            if removed_count > 0:
                self.log.emit(f"- 总共成功删除 {removed_count} 个分节符")
            else:
                self.log.emit("- 未发现或无法删除分节符")
                
        except Exception as e:
            self.log.emit(f"- 删除分节符完全失败：{str(e)}")

    def remove_section_breaks_win32com_fallback(self, doc_path):
        """Win32com备用方案 - 当python-docx无法删除时使用"""
        try:
            import win32com.client
            import pythoncom
            
            self.log.emit("- 尝试使用Win32com删除剩余分节符...")
            
            # 初始化COM
            pythoncom.CoInitialize()
            
            word = win32com.client.DispatchEx("Word.Application")
            word.Visible = False
            word.DisplayAlerts = False
            
            doc = word.Documents.Open(doc_path)
            
            # 使用Word的查找替换删除分节符
            find_obj = doc.Range().Find
            find_obj.ClearFormatting()
            find_obj.Text = "^b"  # 分节符代码
            find_obj.Replacement.ClearFormatting()
            find_obj.Replacement.Text = ""
            
            # 执行替换
            replaced = find_obj.Execute(Replace=win32com.client.constants.wdReplaceAll)
            
            doc.Save()
            doc.Close(SaveChanges=False)
            word.Quit()
            pythoncom.CoUninitialize()
            
            if replaced:
                self.log.emit(f"- Win32com成功删除了额外的分节符")
            else:
                self.log.emit("- Win32com未找到其他分节符")
                
            return True
            
        except Exception as e:
            self.log.emit(f"- Win32com备用方案失败: {e}")
            try:
                if 'doc' in locals():
                    doc.Close(SaveChanges=False)
                if 'word' in locals():
                    word.Quit()
                pythoncom.CoUninitialize()
            except:
                pass
            return False

    
    def has_special_elements(self, paragraph):
        """检查段落是否包含特殊元素 - 支持新旧图片格式"""
        xml = paragraph._p.xml
        # 同时检测 w:drawing (新格式) 和 w:pict (旧格式VML)
        return any(tag in xml for tag in ['<w:drawing', '<w:pict', '<w:tbl', '<w:br', '<w:sectPr'])
    
    def clean_paragraph_spaces(self, paragraph):
        """深度清理段落中的多余空格 - 完全按照V1.0逻辑"""
        if not paragraph.text:
            return
            
        # 获取所有run及其格式信息
        runs_data = []
        for run in paragraph.runs:
            if run.text:
                runs_data.append({
                    'text': run.text,
                    'bold': run.font.bold,
                    'italic': run.font.italic,
                    'underline': run.font.underline,
                    'color': self.get_run_color(run),
                    'highlight': self.get_run_highlight(run)
                })
        
        if runs_data:
            # 合并所有run的文本进行深度清理
            full_text = "".join(run_data['text'] for run_data in runs_data)
            
            # 深度清理各种空白字符
            cleaned_text = self.deep_clean_text_spaces(full_text)
            
            # 如果文本有变化，重新构建段落
            if cleaned_text != full_text:
                # 清除所有run
                for run in paragraph.runs[::-1]:
                    run.clear()
                
                # 重新添加清理后的文本
                if cleaned_text:
                    new_run = paragraph.add_run(cleaned_text)
                    # 如果原来只有一个run，保持其格式
                    if len(runs_data) == 1:
                        self.restore_run_format(new_run, runs_data[0])
    
    def deep_clean_text_spaces(self, text):
        """深度清理文本中的所有类型空格 - 完全按照V1.0逻辑（完全删除所有空格）"""
        if not text:
            return text
        
        # 第一步：删除各种空白字符（包括普通空格）
        cleaned_text = text
        cleaned_text = cleaned_text.replace(' ', '')          # 普通空格（最重要！）
        cleaned_text = cleaned_text.replace('\t', '')         # 制表符
        cleaned_text = cleaned_text.replace('\r', '')         # 回车符
        cleaned_text = cleaned_text.replace('\n', '')         # 换行符
        cleaned_text = cleaned_text.replace('\x0b', '')       # 垂直制表符
        cleaned_text = cleaned_text.replace('\x0c', '')       # 换页符
        cleaned_text = cleaned_text.replace('　', '')         # 全角空格
        cleaned_text = cleaned_text.replace('\u00a0', '')     # 非断行空格
        cleaned_text = cleaned_text.replace('\u2000', '')     # EN QUAD
        cleaned_text = cleaned_text.replace('\u2001', '')     # EM QUAD
        cleaned_text = cleaned_text.replace('\u2002', '')     # EN SPACE
        cleaned_text = cleaned_text.replace('\u2003', '')     # EM SPACE
        cleaned_text = cleaned_text.replace('\u2004', '')     # THREE-PER-EM SPACE
        cleaned_text = cleaned_text.replace('\u2005', '')     # FOUR-PER-EM SPACE
        cleaned_text = cleaned_text.replace('\u2006', '')     # SIX-PER-EM SPACE
        cleaned_text = cleaned_text.replace('\u2007', '')     # FIGURE SPACE
        cleaned_text = cleaned_text.replace('\u2008', '')     # PUNCTUATION SPACE
        cleaned_text = cleaned_text.replace('\u2009', '')     # THIN SPACE
        cleaned_text = cleaned_text.replace('\u200a', '')     # HAIR SPACE
        cleaned_text = cleaned_text.replace('\u200b', '')     # ZERO WIDTH SPACE
        cleaned_text = cleaned_text.replace('\u2028', '')     # LINE SEPARATOR
        cleaned_text = cleaned_text.replace('\u2029', '')     # PARAGRAPH SEPARATOR
        cleaned_text = cleaned_text.replace('\u202f', '')     # NARROW NO-BREAK SPACE
        cleaned_text = cleaned_text.replace('\u205f', '')     # MEDIUM MATHEMATICAL SPACE
        cleaned_text = cleaned_text.replace('\u3000', '')     # IDEOGRAPHIC SPACE (全角空格)
        cleaned_text = cleaned_text.replace('\ufeff', '')     # ZERO WIDTH NO-BREAK SPACE
        
        # 第二步：使用正则表达式删除所有剩余的空白字符
        cleaned_text = re.sub(r'\s+', '', cleaned_text)
        
        return cleaned_text
    
    def remove_sequence_prefix_asterisk(self, text):
        """
        只删除序号前的星号（用于格式标记），保留内容中的星号
        
        例如：
        - '*④RTU...' -> '④RTU...'  （序号前的星号被删除）
        - '*1、投标人...' -> '1、投标人...'  （序号前的星号被删除）
        - '带星号（"*"）' -> '带星号（"*"）'  （内容中的星号被保留）
        """
        if not text:
            return text
        
        # 匹配行首的序号前星号
        # 格式：可选空白 + *或★或※或＊ + 可选空白 + 序号
        # 序号格式包括：
        # - 数字序号：1、1.、1.1、1.1.1 等（后面可接分隔符或空白或中文字符）
        # - 括号序号：(1)、（1）等
        # - 圆圈数字：①②③④⑤⑥⑦⑧⑨⑩ 等（后面可直接跟中文内容）
        # - 中文数字：一、二、等
        pattern = r'^(\s*)[*★※＊](\s*)((?:[（(]\s*\d+(?:\.\d+)*\s*[）)]|\d+(?:[.．]\d+)*[）)]?[、，．.\s]?|[①②③④⑤⑥⑦⑧⑨⑩⑪⑫⑬⑭⑮⑯⑰⑱⑲⑳]|[一二三四五六七八九十]+[、]))'
        
        result = re.sub(pattern, r'\1\2\3', text)
        return result
    
    def set_page_margins(self, doc):
        """设置页边距 - 完全按照V1.0逻辑"""
        self.log.emit("- 设置页边距...")
        
        section = doc.sections[0]
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.18)
        section.right_margin = Cm(3.18)
        
        self.log.emit("- 页边距设置完成")
    
    def set_grid_alignment_options(self, style):
        """设置网格对齐选项 - 完全按照V1.0逻辑（取消勾选两个网格相关选项）"""
        try:
            # 获取段落属性
            pPr = style._element.pPr
            
            # 取消"如果定义了文档网格，则自动调整右缩进(D)"
            # 添加或设置 adjustRightInd 为 0 (false)
            adjust_right_ind = pPr.find(qn('w:adjustRightInd'))
            if adjust_right_ind is not None:
                pPr.remove(adjust_right_ind)
            
            adjust_right_ind_elem = OxmlElement('w:adjustRightInd')
            adjust_right_ind_elem.set(qn('w:val'), '0')
            pPr.append(adjust_right_ind_elem)
            
            # 取消"如果定义了文档网格，则与网格对齐(W)"
            # 添加或设置 snapToGrid 为 0 (false)
            snap_to_grid = pPr.find(qn('w:snapToGrid'))
            if snap_to_grid is not None:
                pPr.remove(snap_to_grid)
            
            snap_to_grid_elem = OxmlElement('w:snapToGrid')
            snap_to_grid_elem.set(qn('w:val'), '0')
            pPr.append(snap_to_grid_elem)
            
        except Exception as e:
            self.log.emit(f"设置网格对齐选项失败: {e}")
    
    def create_custom_styles(self, doc):
        """创建自定义样式 - 标题样式从标书标题2开始，设置后续段落样式为标书正文"""
        self.log.emit("- 创建自定义样式...")

        styles = doc.styles

        # 先创建标书正文样式（因为标题样式需要引用它作为后续段落样式）
        try:
            if "标书正文" in styles:
                styles._element.remove(styles["标书正文"]._element)
        except:
            pass

        body_style = styles.add_style("标书正文", WD_STYLE_TYPE.PARAGRAPH)
        body_style.base_style = None

        font = body_style.font
        font.name = 'SimSun'
        font.size = Pt(10.5)
        font.bold = False

        # 设置英文字体
        try:
            rFonts = body_style._element.rPr.rFonts
            rFonts.set(qn('w:eastAsia'), 'SimSun')
            rFonts.set(qn('w:ascii'), 'Times New Roman')
            rFonts.set(qn('w:hAnsi'), 'Times New Roman')
            rFonts.set(qn('w:cs'), 'Times New Roman')
        except:
            if body_style._element.rPr is None:
                rPr = OxmlElement('w:rPr')
                body_style._element.insert(0, rPr)

            rFonts = OxmlElement('w:rFonts')
            rFonts.set(qn('w:eastAsia'), 'SimSun')
            rFonts.set(qn('w:ascii'), 'Times New Roman')
            rFonts.set(qn('w:hAnsi'), 'Times New Roman')
            rFonts.set(qn('w:cs'), 'Times New Roman')
            body_style._element.rPr.append(rFonts)

        paragraph_format = body_style.paragraph_format
        paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        paragraph_format.space_before = Pt(0)
        paragraph_format.space_after = Pt(0)
        paragraph_format.first_line_indent = Pt(21)  # 2字符缩进
        paragraph_format.left_indent = Pt(0)
        paragraph_format.right_indent = Pt(0)

        # 设置标书正文的后续段落样式为自身
        body_style.next_paragraph_style = body_style

        # 设置网格对齐选项
        self.set_grid_alignment_options(body_style)

        self.log.emit("- 创建标书正文样式")

        # 标书标题样式 (2-6级) - 从标书标题2开始
        for level in range(2, 7):
            style_name = f"标书标题{level}"
            try:
                # 如果样式已存在，删除它
                if style_name in styles:
                    styles._element.remove(styles[style_name]._element)
            except:
                pass

            # 创建新样式
            style = styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)

            # 设置基础样式
            style.base_style = None  # 不继承任何基础样式

            # 设置后续段落样式为标书正文
            style.next_paragraph_style = body_style

            # 字体设置
            font = style.font
            font.name = 'SimSun'
            font.size = Pt(10.5)  # 五号字
            font.bold = True

            # 设置英文字体
            try:
                rFonts = style._element.rPr.rFonts
                rFonts.set(qn('w:eastAsia'), 'SimSun')
                rFonts.set(qn('w:ascii'), 'Times New Roman')
                rFonts.set(qn('w:hAnsi'), 'Times New Roman')
                rFonts.set(qn('w:cs'), 'Times New Roman')
            except:
                # 如果rPr不存在，创建它
                if style._element.rPr is None:
                    rPr = OxmlElement('w:rPr')
                    style._element.insert(0, rPr)

                rFonts = OxmlElement('w:rFonts')
                rFonts.set(qn('w:eastAsia'), 'SimSun')
                rFonts.set(qn('w:ascii'), 'Times New Roman')
                rFonts.set(qn('w:hAnsi'), 'Times New Roman')
                rFonts.set(qn('w:cs'), 'Times New Roman')
                style._element.rPr.append(rFonts)

            # 段落格式
            paragraph_format = style.paragraph_format
            paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            paragraph_format.space_before = Pt(0)
            paragraph_format.space_after = Pt(0)
            paragraph_format.first_line_indent = Pt(0)
            paragraph_format.left_indent = Pt(0)
            paragraph_format.right_indent = Pt(0)

            # 设置大纲级别 (标书标题2对应大纲级别1，标书标题3对应大纲级别2，以此类推)
            try:
                # 确保pPr元素存在
                if style._element.pPr is None:
                    pPr = OxmlElement('w:pPr')
                    style._element.append(pPr)

                # 添加大纲级别元素
                outline_lvl = OxmlElement('w:outlineLvl')
                outline_lvl.set(qn('w:val'), str(level - 1))  # 标书标题2对应大纲级别1
                style._element.pPr.append(outline_lvl)
            except Exception as e:
                self.log.emit(f"设置大纲级别失败: {e}")

            # 设置网格对齐选项
            self.set_grid_alignment_options(style)

            self.log.emit(f"- 创建标书标题{level}样式（大纲级别{level-1}）")

        # 保存文档以确保样式被保存
        try:
            temp_save_file = os.path.join(os.path.dirname(self.input_file), "temp_style_save.docx")
            doc.save(temp_save_file)
            doc = Document(temp_save_file)
            os.remove(temp_save_file)
            self.log.emit("- 样式保存完成")
        except Exception as e:
            self.log.emit(f"- 样式保存失败：{str(e)}")

        self.log.emit("- 自定义样式创建完成")
        return doc
    
    def get_run_color(self, run):
        """获取run的颜色 - 完全按照V1.0逻辑"""
        try:
            return run.font.color.rgb
        except:
            return None
    
    def get_run_highlight(self, run):
        """获取run的突出显示颜色 - 完全按照V1.0逻辑"""
        try:
            return run.font.highlight_color
        except:
            return None
    
    def restore_run_format(self, run, format_info):
        """恢复run的格式 - 完全按照V1.0逻辑"""
        try:
            if format_info['bold']:
                run.font.bold = format_info['bold']
            if format_info['italic']:
                run.font.italic = format_info['italic']
            if format_info['underline']:
                run.font.underline = format_info['underline']
            if format_info['color']:
                run.font.color.rgb = format_info['color']
            if format_info['highlight']:
                run.font.highlight_color = format_info['highlight']
        except:
            pass
    
    def apply_heading_styles(self, doc):
        """应用标题样式 - 样式从标书标题2开始"""
        self.file_logger.info('[apply_heading_styles] 开始')
        self.log.emit("- 识别序号段落...")
        # 提取编号段落
        num_paragraphs = extract_numbered_paragraphs(doc)
        self.log.emit(f"- 找到 {len(num_paragraphs)} 个编号段落")
        self.file_logger.info(f'[apply_heading_styles] 找到 {len(num_paragraphs)} 个编号段落')
        if not num_paragraphs:
            return doc
        # 分析层级
        layered, _ = analyze_hierarchy(num_paragraphs)
        # 应用样式（简化版本，避免多次win32com调用）
        style_applied = 0
        for idx, prefix, level, is_paren, kind, para in layered:
            if level and 1 <= level <= 5:
                # 检查标题长度（不包括前缀）
                text_without_prefix = para.text[len(prefix):].strip()

                # 严格的标题识别条件：
                # 1. 长度限制：不超过15个字符
                # 2. 结尾不能有任何标点符号
                punctuation_marks = ['。', '！', '？', '；', '：', '，', '、', '.', '!', '?', ';', ':', ',']
                has_punctuation_at_end = any(text_without_prefix.endswith(mark) for mark in punctuation_marks)

                # 严格的标题识别条件
                if len(text_without_prefix) <= 15 and not has_punctuation_at_end:
                    try:
                        # 样式名称映射：层级1->标书标题2, 层级2->标书标题3, ..., 层级5->标书标题6
                        style_level = level + 1
                        target_style = doc.styles[f"标书标题{style_level}"]

                        # 应用样式到段落
                        para.style = target_style

                        # 强制刷新段落格式
                        self.force_apply_style_formatting(para, target_style)

                        # 设置大纲级别（通过OXML）- 使用style_level
                        self.set_outline_level(para, style_level)

                        style_applied += 1
                        self.log.emit(f"- 应用标书标题{style_level}样式到：{para.text[:20]}...")
                    except Exception as e:
                        self.log.emit(f"- 应用样式失败：{str(e)}")

        self.log.emit(f"- 已应用 {style_applied} 个标题样式")
        return doc
    
    def set_outline_level(self, paragraph, level):
        """设置段落的大纲级别 - 完全按照V1.0逻辑"""
        try:
            # 确保pPr元素存在
            if paragraph._element.pPr is None:
                pPr = OxmlElement('w:pPr')
                paragraph._element.append(pPr)
            
            # 删除现有的大纲级别
            existing_outline = paragraph._element.pPr.find(qn('w:outlineLvl'))
            if existing_outline is not None:
                paragraph._element.pPr.remove(existing_outline)
            
            # 添加新的大纲级别元素
            outline_lvl = OxmlElement('w:outlineLvl')
            outline_lvl.set(qn('w:val'), str(level - 1))  # 大纲级别从0开始
            paragraph._element.pPr.append(outline_lvl)
        except Exception as e:
            self.log.emit(f"设置大纲级别失败: {e}")
    
    def force_apply_style_formatting(self, paragraph, style):
        """强制应用样式格式到段落 - 完全按照V1.0逻辑"""
        try:
            # 确保段落具有正确的格式
            if style.paragraph_format:
                pf = paragraph.paragraph_format
                sf = style.paragraph_format
                
                # 复制段落格式
                if sf.alignment is not None:
                    pf.alignment = sf.alignment
                if sf.line_spacing_rule is not None:
                    pf.line_spacing_rule = sf.line_spacing_rule
                if sf.space_before is not None:
                    pf.space_before = sf.space_before
                if sf.space_after is not None:
                    pf.space_after = sf.space_after
                if sf.first_line_indent is not None:
                    pf.first_line_indent = sf.first_line_indent
                if sf.left_indent is not None:
                    pf.left_indent = sf.left_indent
                if sf.right_indent is not None:
                    pf.right_indent = sf.right_indent
                              
            
            # 强制应用字体格式到所有run
            self.force_apply_font_formatting(paragraph, style)
                        
        except Exception as e:
            self.log.emit(f"- 强制应用样式格式失败：{str(e)}")
    
    def force_apply_font_formatting(self, paragraph, style):
        """智能应用字体格式，保留颜色、突出显示和局部加粗，并清理空格 - 完全按照V1.0逻辑"""
        try:
            if not paragraph.runs:
                return
            
            # 先进行空格清理
            self.clean_paragraph_spaces(paragraph)
            
            # 遍历所有run，只调整字体族和基础字体，保留其他格式
            for run in paragraph.runs:
                if not run.text.strip():  # 跳过空run
                    continue
                
                # 清理run中的空格
                if run.text:
                    cleaned_text = self.deep_clean_text_spaces(run.text)
                    if cleaned_text != run.text:
                        run.text = cleaned_text
                
                # 保存原有的格式属性
                original_bold = run.font.bold
                original_color = None
                original_highlight = None
                original_underline = run.font.underline
                original_italic = run.font.italic
                
                # 保存颜色
                try:
                    if run.font.color.rgb is not None:
                        original_color = run.font.color.rgb
                except:
                    pass
                
                # 保存突出显示
                try:
                    if hasattr(run.font, 'highlight_color') and run.font.highlight_color is not None:
                        original_highlight = run.font.highlight_color
                except:
                    pass
                
                # 应用基础字体格式（只有当前run没有特殊加粗时才应用样式加粗）
                if style.font:
                    if style.font.name:
                        run.font.name = style.font.name
                    if style.font.size:
                        run.font.size = style.font.size
                    
                    # 只有当前run不是加粗时，才应用样式的加粗设置
                    if original_bold is not True and style.font.bold is not None:
                        run.font.bold = style.font.bold
                
                # 设置字体族（保持原有格式）
                self.set_font_family_preserve_formatting(run, style.name)
                
                # 恢复原有的特殊格式
                if original_color:
                    try:
                        run.font.color.rgb = original_color
                    except:
                        pass
                
                if original_highlight:
                    try:
                        run.font.highlight_color = original_highlight
                    except:
                        pass
                
                # 恢复其他格式
                if original_underline:
                    run.font.underline = original_underline
                if original_italic:
                    run.font.italic = original_italic
            
        except Exception as e:
            self.log.emit(f"- 智能应用字体格式失败：{str(e)}")
    
    def set_font_family_preserve_formatting(self, run, style_name):
        """设置字体族，但保留其他格式 - 完全按照V1.0逻辑"""
        try:
            # 确保rPr元素存在
            if run._element.rPr is None:
                rPr = OxmlElement('w:rPr')
                run._element.insert(0, rPr)
            
            # 只更新字体族，不删除其他格式
            existing_fonts = run._element.rPr.find(qn('w:rFonts'))
            if existing_fonts is not None:
                run._element.rPr.remove(existing_fonts)
            
            # 创建新的字体设置
            rFonts = OxmlElement('w:rFonts')
            
            if style_name.startswith('标书标题') or style_name == '标书正文':
                # 中文字体：宋体，英文字体：Times New Roman
                rFonts.set(qn('w:eastAsia'), 'SimSun')        # 中文字体
                rFonts.set(qn('w:ascii'), 'Times New Roman')  # 英文字体
                rFonts.set(qn('w:hAnsi'), 'Times New Roman')  # 西文字体
                rFonts.set(qn('w:cs'), 'Times New Roman')     # 复杂脚本字体
            
            run._element.rPr.append(rFonts)
            
        except Exception as e:
            self.log.emit(f"设置字体族失败: {e}")
    
    def apply_body_styles(self, doc):
        """应用正文样式 - 完全按照V1.0逻辑"""
        self.file_logger.info('[apply_body_styles] 开始')
        self.log.emit("- 应用正文样式...")

        # 获取所有已经设置为标题样式的段落索引
        heading_indices = set()
        for idx, para in enumerate(doc.paragraphs):
            if para.style.name.startswith('标书标题'):
                heading_indices.add(idx)
        self.file_logger.info(f'[apply_body_styles] 标题段落数: {len(heading_indices)}')
        
        # 为非标题段落应用正文样式
        body_style_applied = 0
        for idx, para in enumerate(doc.paragraphs):
            if idx not in heading_indices and para.text.strip():
                # 检查是否包含特殊元素
                if not self.has_special_elements(para):
                    # 统一应用标书正文样式
                    try:
                        # 获取样式
                        target_style = doc.styles["标书正文"]
                        
                        # 应用样式到段落
                        para.style = target_style
                        
                        # 强制刷新段落格式
                        self.force_apply_style_formatting(para, target_style)
                        
                        body_style_applied += 1
                        
                    except Exception as e:
                        self.log.emit(f"- 应用正文样式失败：{str(e)}")
        
        self.log.emit(f"- 已应用 {body_style_applied} 个正文样式")
        return doc
    
    def format_tables(self, doc):
        """格式化表格 - 完全按照V1.0逻辑"""
        self.file_logger.info('[format_tables] 开始')
        tables = doc.tables
        self.log.emit(f"- 找到 {len(tables)} 个表格")
        self.file_logger.info(f'[format_tables] 表格数量: {len(tables)}')
        
        for table in tables:
            # 设置表格属性 - 根据窗口调整表格
            try:
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                
                # 设置表格自动调整为根据窗口调整
                tbl = table._tbl
                tblPr = tbl.tblPr
                if tblPr is None:
                    tblPr = OxmlElement('w:tblPr')
                    tbl.insert(0, tblPr)
                
                # 删除现有的布局设置
                existing_layout = tblPr.find(qn('w:tblLayout'))
                if existing_layout is not None:
                    tblPr.remove(existing_layout)
                
                # 设置表格布局为自动调整
                tblLayout = OxmlElement('w:tblLayout')
                tblLayout.set(qn('w:type'), 'autofit')
                tblPr.append(tblLayout)
                
                # 设置表格宽度为100%
                existing_width = tblPr.find(qn('w:tblW'))
                if existing_width is not None:
                    tblPr.remove(existing_width)
                
                tblW = OxmlElement('w:tblW')
                tblW.set(qn('w:w'), '5000')  # 100% 宽度
                tblW.set(qn('w:type'), 'pct')
                tblPr.append(tblW)

                # === 新增：设置表格所有边框 ===
                # 清除现有边框
                existing_borders = tblPr.find(qn('w:tblBorders'))
                if existing_borders is not None:
                    tblPr.remove(existing_borders)

                # 创建新边框
                tblBorders = OxmlElement('w:tblBorders')
                border_types = ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']
                for border_type in border_types:
                    border = OxmlElement(f'w:{border_type}')
                    border.set(qn('w:val'), 'single')  # 单线边框
                    border.set(qn('w:sz'), '4')       # 0.5磅 (4 = 0.5 * 8)
                    border.set(qn('w:color'), '000000') # 黑色
                    border.set(qn('w:space'), '0')     # 无间距
                    tblBorders.append(border)

                tblPr.append(tblBorders)
                self.log.emit(f"--- 边框设置完成: {len(border_types)} 种边框类型")
                # === 边框设置结束 ===
                
                self.log.emit(f"- 表格自动调整设置完成")
                
                    
            except Exception as e:
                self.log.emit(f"- 设置表格属性失败：{str(e)}")
            
            # 处理表格中的自动编号
            self.convert_table_auto_numbering(table)
            
            # 遍历所有单元格
            for row_idx, row in enumerate(table.rows):
                for cell in row.cells:
                    # 设置单元格垂直对齐
                    try:
                        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    except:
                        pass

                    # === 新增：设置单元格边框 ===
                    try:
                        tc = cell._tc
                        tcPr = tc.get_or_add_tcPr()
                        
                        # 清除单元格现有边框
                        existing_borders = tcPr.find(qn('w:tcBorders'))
                        if existing_borders is not None:
                            tcPr.remove(existing_borders)
                        
                        # 创建单元格边框元素
                        tcBorders = OxmlElement('w:tcBorders')
                        border_types = ['top', 'left', 'bottom', 'right']
                        for border_type in border_types:
                            border = OxmlElement(f'w:{border_type}')
                            border.set(qn('w:val'), 'single')
                            border.set(qn('w:sz'), '4')
                            border.set(qn('w:color'), '000000')
                            border.set(qn('w:space'), '0')
                            tcBorders.append(border)
                        
                        tcPr.append(tcBorders)
                    except Exception as e:
                        self.log.emit(f"- 设置单元格边框失败：{str(e)}")
                    # === 单元格边框设置结束 ===
                    
                    # 清理和格式化单元格中的段落
                    self.clean_and_format_cell_paragraphs(cell, row_idx == 0)

                    # 确保所有段落都应用了网格对齐选项
                    for para in cell.paragraphs:
                        self.set_table_grid_alignment_options(para)
            
            # 调整表格行高
            self.adjust_table_row_height(table)
        
        self.log.emit(f"- 表格格式化完成")
    
    def convert_table_auto_numbering(self, table):
        """转换表格中的自动编号为文本 - 完全按照V1.0逻辑"""
        try:
            # 使用win32com处理表格中的自动编号
            word = win32com.client.Dispatch("Word.Application")
            word.Visible = False
            
            # 先保存一个临时文件
            temp_file = os.path.join(os.path.dirname(self.input_file), "temp_table_numbering.docx")
            # 这里需要保存整个文档，因为我们只能通过win32com处理整个文档
            # 实际实现时可能需要更复杂的逻辑
            
        except Exception as e:
            self.log.emit(f"- 表格自动编号转换失败：{str(e)}")
    
    def adjust_table_row_height(self, table):
        """根据内容调整表格行高 - 完全按照V1.0逻辑"""
        try:
            for row_idx, row in enumerate(table.rows):
                # 计算该行的最大内容长度
                max_content_length = 0
                has_long_content = False
                
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        # 计算内容长度（中文字符按2个字符计算）
                        content_length = self.calculate_text_display_length(cell_text)
                        max_content_length = max(max_content_length, content_length)
                        
                        # 检查是否有长内容或多行内容
                        if content_length > 20 or '\n' in cell_text:
                            has_long_content = True
                
                # 根据内容调整行高
                try:
                    if max_content_length == 0:
                        # 空行，设置最小行高
                        row.height = Cm(0.5)
                    elif max_content_length <= 10:
                        # 短内容，标准行高
                        row.height = Cm(0.8)
                    elif max_content_length <= 30:
                        # 中等内容，稍高
                        row.height = Cm(1.0)
                    elif has_long_content or max_content_length > 50:
                        # 长内容，较高行高
                        row.height = Cm(1.5)
                    else:
                        # 默认适中行高
                        row.height = Cm(1.2)
                
                except Exception as row_error:
                    # 如果设置行高失败，尝试设置单元格高度
                    try:
                        for cell in row.cells:
                            if max_content_length > 30:
                                cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
                            else:
                                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    except:
                        pass
                        
        except Exception as e:
            self.log.emit(f"- 调整表格行高失败：{str(e)}")
    
    def calculate_text_display_length(self, text):
        """计算文本显示长度（中文字符按2个字符计算） - 完全按照V1.0逻辑"""
        if not text:
            return 0
        
        length = 0
        for char in text:
            # 中文字符、全角字符按2个字符计算
            if ord(char) > 127:
                length += 2
            else:
                length += 1
        return length
    
    def clean_and_format_cell_paragraphs(self, cell, is_header):
        """清理和格式化单元格中的段落，保持原有段落结构"""
        try:
            # 遍历所有段落进行清理和格式化，但不合并
            for para in cell.paragraphs:
                # 清理段落内容
                self.clean_table_paragraph(para)

                # 格式化段落（无论是否有内容）
                self.format_table_paragraph(para, is_header)

                # 设置网格对齐选项
                self.set_table_grid_alignment_options(para)

        except Exception as e:
            self.log.emit(f"- 清理单元格段落失败：{str(e)}")
    
    def clean_table_paragraph(self, paragraph):
        """深度清理表格段落内容 - 完全按照V1.0逻辑"""
        if not paragraph.text:
            return
        
        # 获取原始文本
        original_text = paragraph.text
        
        # 使用统一的深度清理函数
        cleaned_text = self.deep_clean_text_spaces(original_text)
        
        # 如果内容有变化，更新段落
        if cleaned_text != original_text:
            # 保存原有的格式属性（颜色、加粗等）
            original_runs_format = []
            for run in paragraph.runs:
                if run.text:
                    original_runs_format.append({
                        'text': run.text,
                        'bold': run.font.bold,
                        'italic': run.font.italic,
                        'underline': run.font.underline,
                        'color': self.get_run_color(run),
                        'highlight': self.get_run_highlight(run)
                    })
            
            # 清除所有run
            for run in paragraph.runs[::-1]:
                run.clear()
            
            # 如果有清理后的文本，重新创建run
            if cleaned_text:
                new_run = paragraph.add_run(cleaned_text)
                # 如果原来只有一个run，保持其格式
                if len(original_runs_format) == 1:
                    self.restore_run_format(new_run, original_runs_format[0])
    
    def format_table_paragraph(self, paragraph, is_header):
        """格式化表格段落 - 完全按照V1.0逻辑 - 有改进"""
        try:
            # 首先清除段落样式，避免继承"标书正文"等样式的首行缩进
            try:
                # 尝试使用Normal样式，如果存在的话
                paragraph.style = 'Normal'
            except:
                pass
            
            # 设置段落格式
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_before = Pt(0) #段前0字符
            paragraph.paragraph_format.space_after = Pt(0) #段后0字符
            paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE #多倍行距
            paragraph.paragraph_format.line_spacing = 1.15 #1.15倍行距
            paragraph.paragraph_format.left_indent = Pt(0)        # 文本之前0字符
            paragraph.paragraph_format.right_indent = Pt(0)  # 文本之后0字符
            paragraph.paragraph_format.first_line_indent = Pt(0) #首行缩进0字符

            # 设置网格对齐选项（新增）
            self.set_table_grid_alignment_options(paragraph)

            # 强制设置字体格式
            self.force_apply_table_font_formatting(paragraph, is_header)
            
        except Exception as e:
            self.log.emit(f"- 格式化表格段落失败：{str(e)}")
    
    def set_table_grid_alignment_options(self, paragraph):
        """设置表格段落的网格对齐选项 - 新增方法"""
        try:
            # 确保段落有pPr元素
            if paragraph._element.pPr is None:
                pPr = OxmlElement('w:pPr')
                paragraph._element.insert(0, pPr)
            else:
                pPr = paragraph._element.pPr
            
            # 取消"如果定义了文档网格，则自动调整右缩进(D)"
            adjust_right_ind = pPr.find(qn('w:adjustRightInd'))
            if adjust_right_ind is not None:
                pPr.remove(adjust_right_ind)
            
            adjust_right_ind_elem = OxmlElement('w:adjustRightInd')
            adjust_right_ind_elem.set(qn('w:val'), '0')
            pPr.append(adjust_right_ind_elem)
            
            # 取消"如果定义了文档网格，则与网格对齐(W)"
            snap_to_grid = pPr.find(qn('w:snapToGrid'))
            if snap_to_grid is not None:
                pPr.remove(snap_to_grid)
            
            snap_to_grid_elem = OxmlElement('w:snapToGrid')
            snap_to_grid_elem.set(qn('w:val'), '0')
            pPr.append(snap_to_grid_elem)

            # 强制清除所有缩进（移除现有ind元素，创建全新的）
            existing_ind = pPr.find(qn('w:ind'))
            if existing_ind is not None:
                pPr.remove(existing_ind)
            
            # 创建新的ind元素，明确设置所有缩进为0
            ind_elem = OxmlElement('w:ind')
            ind_elem.set(qn('w:firstLine'), '0')  # 首行缩进0
            ind_elem.set(qn('w:left'), '0')       # 左缩进0
            ind_elem.set(qn('w:right'), '0')      # 右缩进0
            pPr.append(ind_elem)

        except Exception as e:
            self.log.emit(f"设置表格网格对齐选项失败: {e}")
    
    def force_apply_table_font_formatting(self, paragraph, is_header):
        """智能应用表格字体格式，保留原有颜色和特殊格式，并清理空格 - 完全按照V1.0逻辑"""
        try:
            if not paragraph.runs:
                return
            
            # 遍历所有run，保留原有的特殊格式
            for run in paragraph.runs:
                if not run.text.strip():  # 跳过空run
                    continue
                
                # 清理run中的空格
                if run.text:
                    cleaned_text = self.deep_clean_text_spaces(run.text)
                    if cleaned_text != run.text:
                        run.text = cleaned_text
                
                # 保存原有的格式属性
                original_bold = run.font.bold
                original_color = None
                original_highlight = None
                original_underline = run.font.underline
                original_italic = run.font.italic
                
                # 保存颜色
                try:
                    if run.font.color.rgb is not None:
                        original_color = run.font.color.rgb
                except:
                    pass
                
                # 保存突出显示
                try:
                    if hasattr(run.font, 'highlight_color') and run.font.highlight_color is not None:
                        original_highlight = run.font.highlight_color
                except:
                    pass
                
                # 设置基础字体格式
                run.font.name = 'SimSun'
                run.font.size = Pt(10.5)
                
                # 表头加粗，但保留原有的加粗格式
                if is_header:
                    run.font.bold = True
                elif original_bold is not True:  # 如果原来不是加粗，则设为不加粗
                    run.font.bold = False
                # 如果原来是加粗，保持加粗
                
                # 设置字体族
                self.set_table_font_family(run)
                
                # 恢复原有的特殊格式
                if original_color:
                    try:
                        run.font.color.rgb = original_color
                    except:
                        pass
                
                if original_highlight:
                    try:
                        run.font.highlight_color = original_highlight
                    except:
                        pass
                
                # 恢复其他格式
                if original_underline:
                    run.font.underline = original_underline
                if original_italic:
                    run.font.italic = original_italic
                
        except Exception as e:
            self.log.emit(f"- 智能应用表格字体格式失败：{str(e)}")
    
    def set_table_font_family(self, run):
        """设置表格字体族 - 完全按照V1.0逻辑"""
        try:
            # 确保rPr元素存在
            if run._element.rPr is None:
                rPr = OxmlElement('w:rPr')
                run._element.insert(0, rPr)
            
            # 只更新字体族
            existing_fonts = run._element.rPr.find(qn('w:rFonts'))
            if existing_fonts is not None:
                run._element.rPr.remove(existing_fonts)
            
            # 创建新的字体设置
            rFonts = OxmlElement('w:rFonts')
            rFonts.set(qn('w:eastAsia'), 'SimSun')        # 中文字体
            rFonts.set(qn('w:ascii'), 'Times New Roman')  # 英文字体
            rFonts.set(qn('w:hAnsi'), 'Times New Roman')  # 西文字体
            rFonts.set(qn('w:cs'), 'Times New Roman')     # 复杂脚本字体
            
            run._element.rPr.append(rFonts)
            
        except Exception as font_error:
            self.log.emit(f"设置表格字体族失败: {font_error}")
    
    def format_images(self, doc):
        """格式化图片 - 完全按照V1.0逻辑"""
        image_count = 0

        for i, para in enumerate(doc.paragraphs):
            xml = para._p.xml
            has_drawing = '<w:drawing' in xml
            has_pict = '<w:pict' in xml

            # 检查段落是否包含图片
            if has_drawing or has_pict:
                image_count += 1
                self.file_logger.debug(f'格式化图片段落{i}: drawing={has_drawing}, pict={has_pict}')

                # 设置���落居中对齐
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                para.paragraph_format.first_line_indent = Pt(0)
                para.paragraph_format.space_before = Pt(0)
                para.paragraph_format.space_after = Pt(0)

                # 调整图片大小
                self.resize_images_in_paragraph(para)

        self.log.emit(f"- 找到并格式化 {image_count} 个图片段落")
    
    def paragraph_has_image(self, paragraph):
        """检查段落是否包含图片 - 支持新旧两种格式"""
        xml = paragraph._p.xml
        return '<w:drawing' in xml or '<w:pict' in xml
    
    def resize_images_in_paragraph(self, paragraph):
        """调整段落中图片的大小"""
        # 页面宽度（考虑页边距，再留0.5cm余量）
        page_width = Inches(8.27) - Cm(3.18) - Cm(3.18) - Cm(0.5)
        page_width_pt = float(page_width / Pt(1))

        for run in paragraph.runs:
            # 处理新格式图片 (w:drawing)
            for drawing in run._element.xpath('.//w:drawing'):
                try:
                    inline = drawing.find('.//wp:inline', drawing.nsmap)
                    if inline is not None:
                        extent = inline.find('.//wp:extent', inline.nsmap)
                        if extent is not None:
                            original_width = int(extent.get('cx'))
                            original_height = int(extent.get('cy'))
                            original_width_pt = original_width / 914400 * 72

                            if original_width_pt > page_width_pt:
                                scale_ratio = page_width_pt / original_width_pt
                                new_width = int(original_width * scale_ratio)
                                new_height = int(original_height * scale_ratio)
                                extent.set('cx', str(new_width))
                                extent.set('cy', str(new_height))
                                self.log.emit(f"- 调整图片大小：原{original_width_pt:.1f}pt -> 新{new_width/914400*72:.1f}pt")
                except Exception as e:
                    pass

            # 处理旧格式VML图片 (w:pict)
            for pict in run._element.xpath('.//w:pict'):
                try:
                    # VML图片尺寸在v:shape的style属性中
                    for shape in pict.xpath('.//*[local-name()="shape"]'):
                        style = shape.get('style', '')
                        width_match = re.search(r'width:\s*([\d.]+)pt', style)
                        height_match = re.search(r'height:\s*([\d.]+)pt', style)
                        if width_match and height_match:
                            original_width_pt = float(width_match.group(1))
                            original_height_pt = float(height_match.group(1))
                            if original_width_pt > page_width_pt:
                                scale_ratio = page_width_pt / original_width_pt
                                new_width_pt = page_width_pt
                                new_height_pt = original_height_pt * scale_ratio
                                new_style = re.sub(r'width:\s*[\d.]+pt', f'width:{new_width_pt:.1f}pt', style)
                                new_style = re.sub(r'height:\s*[\d.]+pt', f'height:{new_height_pt:.1f}pt', new_style)
                                shape.set('style', new_style)
                                self.log.emit(f"- 调整VML图片大小：原{original_width_pt:.1f}pt -> 新{new_width_pt:.1f}pt")
                except Exception as e:
                    pass
    
    def final_cleanup_all_spaces(self, doc):
        """最终全文空格清理 - 完全按照V1.0逻辑（确保所有段落和表格中的空格都被清理）"""
        self.file_logger.info('[final_cleanup_all_spaces] 开始')
        self.log.emit("- 开始最终全文空格清理...")

        # 清理所有段落
        cleaned_paragraphs = 0
        for para in doc.paragraphs:
            if para.text.strip():  # 只处理有内容的段落
                original_text = para.text
                # 直接清理所有run中的空格
                for run in para.runs:
                    if run.text:
                        cleaned_text = self.deep_clean_text_spaces(run.text)
                        # 清理序号前的星号（只删除格式标记星号，保留内容中的星号）
                        cleaned_text = self.remove_sequence_prefix_asterisk(cleaned_text)
                        if cleaned_text != run.text:
                            run.text = cleaned_text
                            cleaned_paragraphs += 1

        # 清理所有表格
        cleaned_tables = 0
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if para.text.strip():  # 只处理有内容的段落
                            for run in para.runs:
                                if run.text:
                                    cleaned_text = self.deep_clean_text_spaces(run.text)
                                    # 清理序号前的星号（只删除格式标记星号，保留内容中的星号）
                                    cleaned_text = self.remove_sequence_prefix_asterisk(cleaned_text)
                                    if cleaned_text != run.text:
                                        run.text = cleaned_text
                                        cleaned_tables += 1

        self.log.emit(f"- 最终空格清理完成：处理 {cleaned_paragraphs} 个段落，{cleaned_tables} 个表格单元格")
    
    def highlight_keywords(self, doc):
        """关键词标青绿色处理 + 符号标红处理"""
        self.log.emit("- 开始关键词标青绿色和符号标红处理...")
        
        # 定义需要标青绿色的关键词
        highlight_keywords = ["提供", "出具", "具有", "详细说明", "详细描述", "检验报告"]
        
        # 创建关键词的正则表达式模式
        pattern = '(' + '|'.join(re.escape(keyword) for keyword in highlight_keywords) + ')'

            # 定义需要标红的符号模式
        symbol_pattern = r'([*★※＊])'
    
        keyword_highlighted_count = 0 # 关键词计数
        symbol_highlighted_count = 0 # 符号计数
        
        # 处理所有段落
        for para in doc.paragraphs:
            if not para.text.strip():
                continue
                
            # 检查段落是否包含关键词
            if any(keyword in para.text for keyword in highlight_keywords):
                keyword_highlighted_count += self.highlight_keywords_in_paragraph(para, pattern, highlight_keywords)
            
            # 检查段落是否包含特殊符号（只处理序号前后的符号）
            # 匹配行首的序号标记：*1、*1.、1、*、1.*、1.1. *、*(1)、1）*等模式
            sequence_symbol_pattern = r'^\s*([*★※＊]\s*)?(?:[（(]\s*\d+(?:\.\d+)*\s*[）)]|\d+(?:\.\d+)*[）)]?)[、，\.\s]*([*★※＊]\s*)?'
            if re.search(sequence_symbol_pattern, para.text):
                symbol_highlighted_count += self.highlight_symbols_in_paragraph(para, symbol_pattern)
        
        # 处理表格中的内容
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if not para.text.strip():
                            continue
                        if any(keyword in para.text for keyword in highlight_keywords):
                            keyword_highlighted_count += self.highlight_keywords_in_paragraph(para, pattern, highlight_keywords)
                        if re.search(sequence_symbol_pattern, para.text):
                            symbol_highlighted_count += self.highlight_symbols_in_paragraph(para, symbol_pattern)
        
        self.log.emit(f"- 关键词标青绿色处理完成，处理了 {keyword_highlighted_count} 个关键词")
        self.log.emit(f"- 符号标红处理完成，处理了 {symbol_highlighted_count} 个符号")

    def highlight_keywords_with_options(self, doc, keyword_enabled=True, symbol_enabled=True):
        """带选项的关键词和符号处理"""
        self.log.emit("- 开始标记处理...")

        # 定义需要标青绿色的关键词
        highlight_keywords = ["提供", "出具", "具有", "详细说明", "详细描述", "检验报告"]

        # 创建关键词的正则表达式模式
        pattern = '(' + '|'.join(re.escape(keyword) for keyword in highlight_keywords) + ')'

        # 定义需要标红的符号模式
        symbol_pattern = r'([*★※＊])'

        # 匹配行首的序号标记
        sequence_symbol_pattern = r'^\s*([*★※＊]\s*)?(?:[（(]\s*\d+(?:\.\d+)*\s*[）)]|\d+(?:\.\d+)*[）)]?)[、，\.\s]*([*★※＊]\s*)?'

        keyword_highlighted_count = 0
        symbol_highlighted_count = 0

        # 处理所有段落
        for para in doc.paragraphs:
            if not para.text.strip():
                continue

            # 处理关键词
            if keyword_enabled and any(keyword in para.text for keyword in highlight_keywords):
                keyword_highlighted_count += self.highlight_keywords_in_paragraph(para, pattern, highlight_keywords)

            # 处理符号
            if symbol_enabled and re.search(sequence_symbol_pattern, para.text):
                symbol_highlighted_count += self.highlight_symbols_in_paragraph(para, symbol_pattern)

        # 处理表格中的内容
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if not para.text.strip():
                            continue
                        if keyword_enabled and any(keyword in para.text for keyword in highlight_keywords):
                            keyword_highlighted_count += self.highlight_keywords_in_paragraph(para, pattern, highlight_keywords)
                        if symbol_enabled and re.search(sequence_symbol_pattern, para.text):
                            symbol_highlighted_count += self.highlight_symbols_in_paragraph(para, symbol_pattern)

        if keyword_enabled:
            self.log.emit(f"- 关键词标青绿色处理完成，处理了 {keyword_highlighted_count} 个关键词")
        if symbol_enabled:
            self.log.emit(f"- 符号标红处理完成，处理了 {symbol_highlighted_count} 个符号")

    def remove_headers_and_footers(self, doc):
        """删除文档中的页眉和页脚"""
        self.file_logger.info('[remove_headers_and_footers] 开始')
        self.log.emit("- 开始删除页眉页脚...")

        removed_count = 0
        try:
            for section in doc.sections:
                # 删除页眉 - 清空所有内容（包括图片等）
                header = section.header
                if header is not None:
                    for para in header.paragraphs:
                        p = para._element
                        p.getparent().remove(p)
                        removed_count += 1

                # 删除页脚 - 清空所有内容（包括图片等）
                footer = section.footer
                if footer is not None:
                    for para in footer.paragraphs:
                        p = para._element
                        p.getparent().remove(p)
                        removed_count += 1

                # 处理首页页眉页脚
                if section.different_first_page_header_footer:
                    first_header = section.first_page_header
                    if first_header is not None:
                        for para in first_header.paragraphs:
                            p = para._element
                            p.getparent().remove(p)
                            removed_count += 1

                    first_footer = section.first_page_footer
                    if first_footer is not None:
                        for para in first_footer.paragraphs:
                            p = para._element
                            p.getparent().remove(p)
                            removed_count += 1

            if removed_count > 0:
                self.log.emit(f"- 页眉页脚删除完成，共处理 {removed_count} 个")
            else:
                self.log.emit("- 未发现需要删除的页眉页脚内容")

        except Exception as e:
            self.log.emit(f"- 删除页眉页脚时出错: {str(e)}")

    def highlight_keywords_in_paragraph(self, paragraph, pattern, highlight_keywords):
        """在段落中标青绿色关键词 - 完全按照V1.0逻辑"""
        try:
            if not paragraph.text.strip():
                return 0
            
            # 按关键词分割文本
            parts = re.split(pattern, paragraph.text)
            
            if len(parts) <= 1:  # 没有关键词
                return 0
            
            # 保存原始段落的格式信息
            original_alignment = paragraph.alignment
            original_pf = paragraph.paragraph_format
            
            # 清除所有run
            for run in paragraph.runs[::-1]:
                run.clear()
            
            highlighted_count = 0
            
            # 重新构建段落
            for part in parts:
                if not part:
                    continue
                
                # 判断是否为关键词
                is_keyword = part in highlight_keywords
                
                if is_keyword:
                    highlighted_count += 1
                
                # 按字符类型处理每个部分
                current_run = None
                current_type = None
                
                for ch in part:
                    ch_type = 'en' if ord(ch) < 128 else 'cn'
                    if ch_type != current_type:
                        current_run = paragraph.add_run()
                        
                        # 如果是关键词，设置青绿色底纹
                        if is_keyword:
                            try:
                                # 设置青绿色底纹（突出显示）
                                current_run.font.highlight_color = 3  # 青绿色底纹
                                # 文字保持黑色
                                current_run.font.color.rgb = None  # 保持默认黑色
                            except Exception as e:
                                self.log.emit(f"设置关键词底纹失败: {e}")
                        
                        # 设置字体
                        if ch_type == 'en':
                            current_run.font.name = 'Times New Roman'
                            current_run.font.size = Pt(10.5)
                            try:
                                if current_run._element.rPr is not None:
                                    current_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
                            except:
                                pass
                        else:
                            current_run.font.name = 'SimSun'
                            current_run.font.size = Pt(10.5)
                            try:
                                if current_run._element.rPr is not None:
                                    current_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
                            except:
                                pass
                        current_type = ch_type
                    
                    if current_run is not None:
                        current_run.text += ch
            
            # 恢复段落格式
            try:
                paragraph.alignment = original_alignment
            except:
                pass
            
            return highlighted_count
            
        except Exception as e:
            self.log.emit(f"- 段落关键词标青绿色失败：{str(e)}")
            return 0
    
    def highlight_symbols_in_paragraph(self, paragraph, pattern):
        """在段落中标红序号前后的特殊符号（不标红段落中间的符号）"""
        try:
            text = paragraph.text
            if not text.strip():
                return 0

            # 只匹配序号前后的符号位置
            # 匹配模式：行首可选空白 + 可选符号 + 序号 + 可选符号
            seq_pattern = r'^(\s*)([*★※＊])?(\s*)([（(]\s*\d+(?:\.\d+)*\s*[）)]|\d+(?:[.．]\d+)*[）)]?)[、，．.\s]*([*★※＊])?'
            match = re.match(seq_pattern, text)
            if not match:
                return 0

            # 找出需要标红的符号位置
            symbol_positions = set()
            if match.group(2):  # 序号前的符号
                pos = len(match.group(1))  # 空白后的位置
                symbol_positions.add(pos)
            if match.group(5):  # 序号后的符号
                pos = match.end(5) - 1
                symbol_positions.add(pos)

            if not symbol_positions:
                return 0

            # 保存原始段落格式
            original_alignment = paragraph.alignment

            # 清除所有run
            for run in paragraph.runs[::-1]:
                run.clear()

            highlighted_count = 0

            # 逐字符重建段落
            for i, ch in enumerate(text):
                run = paragraph.add_run(ch)
                # 设置字体
                if ord(ch) < 128:
                    run.font.name = 'Times New Roman'
                else:
                    run.font.name = 'SimSun'
                run.font.size = Pt(10.5)

                # 如果是需要标红的符号位置
                if i in symbol_positions and ch in '*★※＊':
                    run.font.highlight_color = 6  # 红色底纹
                    highlighted_count += 1

            # 恢复段落格式
            try:
                paragraph.alignment = original_alignment
            except:
                pass

            return highlighted_count

        except Exception as e:
            return 0
    
    # ==================== 智能应答相关方法 ====================
    
    def execute_response_processing(self, doc, progress_offset=10, progress_range=90):
        """执行完整的应答处理流程 - 完全按照V2.0逻辑"""
        self.file_logger.info('[execute_response_processing] 开始')
        step_size = progress_range / 7

        # 第零阶段：预处理
        self.log.emit("预处理：检查并转换自动编号...")
        self.file_logger.info('[execute_response_processing] 预处理：检查并转换自动编号')
        # 新增：确保自动编号被转换为文本，否则正则无法匹配
        doc = self.convert_auto_numbering_to_text(doc)

        self.log.emit("预处理表格、图片、分页符...")
        self.file_logger.info('[execute_response_processing] 预处理表格、图片、分页符')
        doc = self.insert_empty_paragraph_after_tables_and_drawings(doc)
        self._count_images(doc, '预处理后')
        self.progress.emit(int(progress_offset + step_size * 1))

        # 第一阶段：选出序号
        self.log.emit("识别文档序号...")
        self.file_logger.info('[execute_response_processing] 识别文档序号')
        num_paragraphs = extract_numbered_paragraphs(doc)
        self.log.emit(f"找到 {len(num_paragraphs)} 个序号段落")
        self.file_logger.info(f'[execute_response_processing] 找到 {len(num_paragraphs)} 个序号段落')
        self.progress.emit(int(progress_offset + step_size * 2))

        # 第二阶段：动态排序
        self.log.emit("对序号进行层级分析...")
        self.file_logger.info('[execute_response_processing] 层级分析')
        layered, headings12 = analyze_hierarchy(num_paragraphs)
        self.file_logger.info(f'[execute_response_processing] 层级分析完成, layered={len(layered)}, headings12={len(headings12)}')
        self.progress.emit(int(progress_offset + step_size * 3))

        # 第三阶段：获取索引
        self.log.emit("整理序号索引...")
        numeric_indices = [item[0] for item in layered if item[2] is not None]
        numeric_indices.sort()
        self.file_logger.info(f'[execute_response_processing] 序号索引数量: {len(numeric_indices)}')
        self.progress.emit(int(progress_offset + step_size * 4))

        # 第四阶段：选出最小层级
        self.log.emit("筛选最小层级条款...")
        minimal_clauses = find_minimal_clauses(layered)
        self.log.emit(f"找到 {len(minimal_clauses)} 个需要应答的条款")
        self.file_logger.info(f'[execute_response_processing] 最小层级条款: {len(minimal_clauses)}')
        self.progress.emit(int(progress_offset + step_size * 5))

        # 第六阶段：插入回应
        self.log.emit("插入应答内容...")
        self.file_logger.info('[execute_response_processing] 开始插入应答')
        self.insert_responses(doc, minimal_clauses, numeric_indices, layered)
        self._count_images(doc, '插入应答后')
        self.progress.emit(int(progress_offset + step_size * 6))

        # 第七阶段：格式完善
        self.log.emit("优化文档格式...")
        self.file_logger.info('[execute_response_processing] 优化文档格式')
        self.remove_only_pure_empty_paragraphs(doc)
        self._count_images(doc, '格式优化后')
        self.progress.emit(int(progress_offset + step_size * 7))
        self.file_logger.info('[execute_response_processing] 完成')
    
    def insert_empty_paragraph_after_tables_and_drawings(self, doc):
        """在所有表格和图片后插入空行"""
        tables = doc.tables
        self.log.emit(f"文档中共有 {len(tables)} 个表格")
        
        processed_positions = set()
        
        if hasattr(doc, '_body') and hasattr(doc._body, '_body'):
            body_elements = list(doc._body._body)
            
            for i in range(len(body_elements)-1, -1, -1):
                elem = body_elements[i]
                
                if elem.tag.endswith('tbl'):
                    para_count = 0
                    for j in range(i):
                        if body_elements[j].tag.endswith('p'):
                            para_count += 1
                    
                    position = para_count
                    
                    if position in processed_positions:
                        continue
                    
                    processed_positions.add(position)
                    
                    next_para = None
                    for j in range(i+1, len(body_elements)):
                        if body_elements[j].tag.endswith('p'):
                            next_para_idx = para_count
                            for k in range(i+1, j):
                                if body_elements[k].tag.endswith('p'):
                                    next_para_idx += 1
                            
                            if next_para_idx < len(doc.paragraphs):
                                next_para = doc.paragraphs[next_para_idx]
                            break
                    
                    if next_para:
                        new_p = OxmlElement('w:p')
                        next_para._p.addprevious(new_p)
                    else:
                        new_p = OxmlElement('w:p')
                        doc._body._body.append(new_p)
        
        # 处理图片和分页符
        for i in range(len(doc.paragraphs)-1, -1, -1):
            para = doc.paragraphs[i]
            
            if hasattr(para, '_p') and para._p is not None:
                if '<w:drawing' in para._p.xml or 'w:br w:type="page"' in para._p.xml or '<w:sectPr' in para._p.xml:
                    if i in processed_positions:
                        continue
                    
                    processed_positions.add(i)
                    
                    if i < len(doc.paragraphs) - 1:
                        next_para = doc.paragraphs[i+1]
                        if not next_para.text.strip():
                            continue
                        
                        new_p = OxmlElement('w:p')
                        next_para._p.addprevious(new_p)
                    else:
                        new_p = OxmlElement('w:p')
                        doc._body._body.append(new_p)
        
        return doc

    def insert_responses(self, doc, minimal_clauses, numeric_indices, layered):
        """插入应答内容"""
        self.file_logger.info(f'[insert_responses] 开始, 条款数={len(minimal_clauses)}')
        self.file_logger.info(f'[insert_responses] 使用产品: {self.selected_product or "常规应答（无指定产品）"}')

        # 记录模板加载状态
        load_all_templates()
        self.file_logger.info(f'[insert_responses] 已加载产品模板: {list(_product_templates.keys())}')
        self.file_logger.info(f'[insert_responses] 默认模板条款数: {len(_default_template)}')
        clause_pattern = re.compile(r'''
            ^\s*                  # 行首可选空白
            [*★※＊]?                 # 可选 * 或 ★ 或 ※ 或 ＊
            \s*
            (
              第[一二三四五六七八九十]+[章节]         |  # 第N章 或 第N节
              [一二三四五六七八九十]+、               |  # 中文"一、二、…"
              [\(\（][一二三四五六七八九十]+[\)\）]、?  |  # （一）或（五）、
              [\(\（]\d+[\)\）]、?                    |  # (1) 或 (2)、
              \d+[）\)]                               |  # 1) 或 1）
              \d+(?:\.\d+)*\.?(?!\s*[%％\/\-])  # 1.、1.1、1.1.1
            )
        ''', re.VERBOSE)

        # 构建layered索引到(idx, lvl, kind)的映射，用于计算条款内容范围
        layered_info = {item[0]: (item[2], item[4]) for item in layered}

        for idx, prefix, lvl, is_paren, kind, clause_para in reversed(minimal_clauses):
            # pos 用于向上查找父条款的扩展匹配逻辑
            pos = numeric_indices.index(idx)

            # 计算条款内容范围：找到下一个同级或更高级的条款
            # 规则：当下一个序号 C 的层级 <= 当前序号 B 的层级时，停止（即 next_lvl <= lvl）
            # 这样可以确保子条款（层级更深）都被包含在当前条款范围内
            next_clause_start = len(doc.paragraphs)
            for next_idx in numeric_indices[pos + 1:]:
                if next_idx in layered_info:
                    next_lvl, next_kind = layered_info[next_idx]
                    # 找到同级或更高级的条款（next_lvl <= lvl），停止
                    if next_lvl <= lvl:
                        next_clause_start = next_idx
                        break

            # end_idx 使用 next_clause_start - 1，确保应答插入在条款内容之后
            end_idx = next_clause_start - 1

            clause_content_lines = []
            for j in range(idx, next_clause_start):
                raw = doc.paragraphs[j].text or ""
                if j == idx:
                    m = clause_pattern.match(raw)
                    if m:
                        prefix = m.group(1)
                        # 验证序号的合理性，比如检查是否符合常见的序号格式
                        if re.match(r'^\d+(\.\d{1,2}){0,3}\.?$', prefix):
                            body = raw[m.end():]
                        else:
                            # 如果序号看起来不合理，可能需要重新处理
                            body = raw
                    body = raw[m.end():] if m else raw
                    body = re.sub(r'^[\s.．\*、\-•※]+', '', body)
                else:
                    body = raw

                # 清理序号前的星号（只删除格式标记星号，保留内容中的星号）
                body = self.remove_sequence_prefix_asterisk(body)

                clause_content_lines.append(body)
            
            # ========== 智能应答核心逻辑（按PRD优先级）==========
            # 原始条款内容（未经process_clauses处理）
            original_clause_lines = clause_content_lines.copy()
            
            # 用于存储完整的应答数据
            response_data = None
            matched = False
            similarity = 0.0
            source = None
            
            # 【优先级1+2】：先检查关键词规则（标题+内容关键词匹配）
            keyword_matched, keyword_response, rule_name = self.check_keyword_rules(original_clause_lines)
            
            if keyword_matched:
                # 命中关键词规则，直接使用规则应答
                template_content = keyword_response
                matched = True
                source = rule_name
                self.log.emit(f"  条款命中关键词规则 [{rule_name}]")
                self.file_logger.info(f'[insert_responses] 条款{idx}命中关键词规则: {rule_name}')
                matched_keyword = rule_name
            else:
                # 【优先级3】：关键词未命中，尝试相似度匹配（V3.0功能）
                # 先对条款进行process_clauses处理
                clause_content_lines = self.process_clauses(original_clause_lines)
                clause_full_text = '\n'.join(clause_content_lines)
                
                # 只有选择了产品才进行相似度匹配
                if self.selected_product:
                    matched, response_data, similarity, source = match_clause_with_product(
                        clause_full_text, self.selected_product
                    )

                    # 如果当前条款未匹配，且条款内容很短（<30字符），尝试向上查找父条款一起匹配
                    clause_text_length = len(clause_full_text.replace('\n', '').replace(' ', ''))
                    if not matched and pos > 0 and clause_text_length < 30:
                        best_extended_match = (matched, response_data, similarity, source)

                        for back_steps in range(1, min(6, pos + 1)):
                            parent_pos = pos - back_steps
                            parent_idx = numeric_indices[parent_pos]

                            extended_lines = []
                            for j in range(parent_idx, next_clause_start):
                                raw = doc.paragraphs[j].text or ""
                                if raw.strip():
                                    extended_lines.append(raw.strip())
                            extended_text = '\n'.join(extended_lines)

                            matched2, response_data2, similarity2, source2 = match_clause_with_product(
                                extended_text, self.selected_product
                            )

                            if matched2 and similarity2 > best_extended_match[2]:
                                best_extended_match = (matched2, response_data2, similarity2, source2)

                        # 要求组合匹配的相似度必须达到更高的标准（0.70），因为这已经是"尝试挽救"的步骤
                        # 避免短条款因凑巧与父条款组合后勉强过线而导致的误匹配
                        if best_extended_match[0] and best_extended_match[2] >= 0.70 and best_extended_match[2] > similarity:
                            matched, response_data, similarity, source = best_extended_match

                # 根据匹配结果确定应答内容
                if matched:
                    # 相似度匹配成功，使用模板应答
                    if isinstance(response_data, dict):
                        response_text = response_data.get('response_text', _default_response)
                    else:
                        response_text = response_data
                    template_content = response_text.split('\n') if isinstance(response_text, str) else [response_text]
                    self.log.emit(f"  条款相似度匹配成功 [相似度:{similarity:.2f}] 来源:{source}")
                    self.file_logger.info(f'[insert_responses] 条款{idx}相似度匹配成功: 相似度={similarity:.2f}, 来源={source}')
                    matched_keyword = source
                else:
                    # 未匹配时的处理逻辑（按PRD区分快速应答和详细应答）
                    if self.selected_product:
                        # 详细应答模式（选择了产品）：使用 process_clauses() 处理后的内容
                        template_content = clause_content_lines if clause_content_lines else [_default_response]
                        self.log.emit(f"  条款未匹配 [最高相似度:{similarity:.2f}]，使用通用规则生成应答")
                        self.file_logger.info(f'[insert_responses] 条款{idx}未匹配(详细应答): 最高相似度={similarity:.2f}, 使用通用规则')
                    else:
                        # 常规应答模式（未选择产品）：使用通用规则生成应答
                        template_content = clause_content_lines if clause_content_lines else []
                        self.log.emit(f"  常规应答模式，使用通用规则生成应答")
                        self.file_logger.info(f'[insert_responses] 条款{idx}常规应答: 使用通用规则')
                    matched_keyword = None

            # 从response_data中获取is_superior标记
            is_superior = False
            if isinstance(response_data, dict):
                is_superior = response_data.get('is_superior', False)

            # 兼容旧方式：检测模板文本中是否包含[优于]标记
            if not is_superior:
                template_text = '\n'.join(template_content)
                if '[优于]' in template_text:
                    is_superior = True
                    # 从模板内容中移除[优于]标记
                    template_content = [line.replace('[优于]', '').strip() for line in template_content]
                    # 移除因清理标记而变成空的行
                    template_content = [line for line in template_content if line]

            # 定位插入点
            paras = doc.paragraphs
            insert_pos = end_idx
            
            def has_page_or_section_break(para):
                if not hasattr(para, '_p') or para._p is None:
                    return False, False
                
                xml_str = para._p.xml
                has_page_break = '<w:br w:type="page"/>' in xml_str
                has_section_break = '<w:sectPr' in xml_str
                
                return has_page_break, has_section_break
            
            last_para = paras[end_idx]
            last_page_break, last_section_break = has_page_or_section_break(last_para)
            
            if last_page_break or last_section_break:
                insert_pos = max(end_idx - 1, 0)
            else:
                insert_pos = end_idx
            
            insert_pos = max(0, min(insert_pos, len(paras) - 1))
            insert_after_para = paras[insert_pos]
            
            # 插入应答
            answer_intro_para = self.insert_paragraph_after(insert_after_para)

            # 设置应答引导段落格式
            answer_intro_para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            answer_intro_para.paragraph_format.first_line_indent = Pt(21)
            answer_intro_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            answer_intro_para.paragraph_format.space_before = Pt(0)
            answer_intro_para.paragraph_format.space_after = Pt(0)
            self.set_grid_alignment_options_paragraph(answer_intro_para)

            # 添加应答引导文本
            intro_run = answer_intro_para.add_run("投标人应答：")
            intro_run.bold = True
            intro_run.font.name = 'SimSun'
            intro_run.font.size = Pt(10.5)
            try:
                if intro_run._element.rPr is not None:
                    intro_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
            except:
                pass

            # 第一部分：标准应答（固定内容，与"投标人应答："在同一行）
            if is_superior:
                self.add_simple_text(answer_intro_para, "我方完全满足且优于招标文件要求。")
                self.log.emit(f"  检测到[优于]标记，已插入优于应答")
            else:
                self.add_simple_text(answer_intro_para, "我方完全满足招标文件要求。")

            # 第二部分：具体应答内容（换行）
            last_insert_point = answer_intro_para

            # 检查是否有完整的模板元素数据（包含表格/图片）
            has_rich_content = (
                isinstance(response_data, dict) and
                response_data.get('element_indices') and
                len(response_data.get('element_indices', [])) > 0
            )

            if has_rich_content and matched:
                # 使用新的元素复制方法，保留表格和图片
                self.log.emit(f"  使用元素复制方式插入应答（包含 {len(response_data.get('element_indices', []))} 个元素）")
                last_para = self.copy_template_elements(doc, last_insert_point, response_data)
            else:
                # 逐行插入文本（从第1行开始，全部作为具体应答内容）
                last_para = last_insert_point
                for i, line in enumerate(template_content):
                    # 每行都创建新段落
                    last_para = self.insert_paragraph_after(last_para)
                    last_para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    last_para.paragraph_format.first_line_indent = Pt(21)
                    last_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
                    last_para.paragraph_format.space_before = Pt(0)
                    last_para.paragraph_format.space_after = Pt(0)

                    # 设置网格对齐选项
                    self.set_grid_alignment_options_paragraph(last_para)

                    # 清理序号前的星号（只删除格式标记星号，保留内容中的星号）
                    clean_line = self.remove_sequence_prefix_asterisk(line)

                    # 添加模板文本
                    self.add_simple_text(last_para, clean_line)

    def copy_clause_tables(self, doc, start_para_idx, end_para_idx, insert_after_para):
        """
        复制输入文档中条款范围内的表格到应答内容后面

        参数：
            doc: 文档对象
            start_para_idx: 条款起始段落索引
            end_para_idx: 条款结束段落索引（不包含）
            insert_after_para: 在此段落后插入表格
        返回：
            最后插入的元素（段落或表格后的段落）
        """
        import copy

        # 获取文档body中的所有元素
        if not hasattr(doc, '_body') or not hasattr(doc._body, '_body'):
            return insert_after_para

        body_elements = list(doc._body._body)

        # 建立段落索引到body元素索引的映射
        para_to_body_idx = {}
        para_count = 0
        for body_idx, element in enumerate(body_elements):
            tag = element.tag.split('}')[-1] if '}' in element.tag else element.tag
            if tag == 'p':
                para_to_body_idx[para_count] = body_idx
                para_count += 1

        # 确定条款在body中的范围
        if start_para_idx not in para_to_body_idx:
            return insert_after_para

        body_start = para_to_body_idx[start_para_idx]

        # 找到结束位置
        if end_para_idx in para_to_body_idx:
            body_end = para_to_body_idx[end_para_idx]
        else:
            # 如果end_para_idx超出范围，使用最后一个元素
            body_end = len(body_elements)

        # 查找范围内的表格
        tables_to_copy = []
        for body_idx in range(body_start, body_end):
            element = body_elements[body_idx]
            tag = element.tag.split('}')[-1] if '}' in element.tag else element.tag
            if tag == 'tbl':
                tables_to_copy.append(element)

        if not tables_to_copy:
            return insert_after_para

        self.log.emit(f"  发现条款中包含 {len(tables_to_copy)} 个表格，复制到应答后")

        # 获取插入位置
        last_element = insert_after_para._p if hasattr(insert_after_para, '_p') else None
        last_para = insert_after_para

        for table_element in tables_to_copy:
            # 深度复制表格
            copied_table = copy.deepcopy(table_element)

            # 在当前位置后插入复制的表格
            if last_element is not None:
                last_element.addnext(copied_table)
                last_element = copied_table

            # 在表格后添加一个空段落，便于后续插入
            empty_para = self.insert_paragraph_after_element(doc, copied_table)
            if empty_para:
                last_element = empty_para._p
                last_para = empty_para

        return last_para

    def insert_paragraph_after_element(self, doc, element):
        """在指定元素后插入一个空段落"""
        from docx.oxml import OxmlElement

        new_p = OxmlElement('w:p')
        element.addnext(new_p)

        # 创建Paragraph对象
        from docx.text.paragraph import Paragraph
        return Paragraph(new_p, doc._body)

    def set_grid_alignment_options_paragraph(self, paragraph):
        """设置段落网格对齐选项"""
        try:
            # 确保段落有pPr元素
            if paragraph._element.pPr is None:
                pPr = OxmlElement('w:pPr')
                paragraph._element.insert(0, pPr)
            else:
                pPr = paragraph._element.pPr
            
            # 取消网格对齐
            adjust_right_ind = pPr.find(qn('w:adjustRightInd'))
            if adjust_right_ind is not None:
                pPr.remove(adjust_right_ind)
            
            adjust_right_ind_elem = OxmlElement('w:adjustRightInd')
            adjust_right_ind_elem.set(qn('w:val'), '0')
            pPr.append(adjust_right_ind_elem)
            
            snap_to_grid = pPr.find(qn('w:snapToGrid'))
            if snap_to_grid is not None:
                pPr.remove(snap_to_grid)
            
            snap_to_grid_elem = OxmlElement('w:snapToGrid')
            snap_to_grid_elem.set(qn('w:val'), '0')
            pPr.append(snap_to_grid_elem)
            
        except Exception as e:
            self.log.emit(f"设置段落网格对齐选项失败: {e}")

    def check_keyword_rules(self, clause_content_lines):
        """
        检查条款是否命中PRD中的优先级1和2关键词规则
        
        返回:
            (matched, response, rule_name)
            - matched: 是否命中关键词规则
            - response: 匹配的应答内容（列表格式）
            - rule_name: 命中的规则名称（用于日志）
        """
        if not clause_content_lines:
            return False, None, None
        
        # 清理空行
        lines = [line for line in clause_content_lines if line.strip()]
        if not lines:
            return False, None, None
        
        first_line = lines[0].strip()
        full_text = '\n'.join(lines)

        # ========== 优先级1：基于标题关键词匹配 ==========
        # 标题判定条件：≤15字符 且 句尾无标点符号
        is_title = len(first_line) <= 15 and not re.search(r'[。！？；：，.!?;:,]$', first_line)

        if is_title:
            # 规则1：项目概况类
            keywords_group1 = ["项目概况", "工程概况", "工程范围", "线路概况", "正线主要技术标准", "联络线技术标准",
                              "招标范围", "适用范围", "使用环境条件", "设计标准", "环境条件", "线路主要技术标准", "设计寿命"]
            for kw in keywords_group1:
                if kw in first_line:
                    return True, [f"我已了解本项目{first_line}，我方所供产品完全满足本项目{first_line}的要求。"], f"标题规则-概况类({kw})"

            # 规则2：环境地质类
            keywords_group2 = ["使用环境", "地貌特征", "工程地质", "水文地质", "地震动参数", "地震", "海拔", "气象特征", "气象条件"]
            for kw in keywords_group2:
                if kw in first_line:
                    return True, [f"我已了解本项目{first_line}，我方所供产品完全适用于本项目{first_line}的情况。"], f"标题规则-环境类({kw})"

            # 规则3：设备清单类
            keywords_group3 = ["设备清单", "招标数量", "供货数量及规格", "需求数量", "需求一览表", "备品备件", "专用工具"]
            for kw in keywords_group3:
                if kw in first_line:
                    return True, [f"我方按照{first_line}要求提供设备，详见商务报价部分。"], f"标题规则-设备类({kw})"

            # 规则4：技术服务
            if "技术服务" in first_line:
                return True, [f"我方将按照招标文件要求提供{first_line}。"], "标题规则-技术服务"

            # 规则5：试验类
            keywords_group5 = ["型式试验", "出厂试验", "现场试验"]
            for kw in keywords_group5:
                if kw in first_line:
                    return True, [f"我方所供设备将按照招标文件的{first_line}的要求来进行试验。"], f"标题规则-试验类({kw})"

            # 规则6：铭牌包装类（仅当条款只有一行时使用简短应答）
            if len(lines) == 1:
                keywords_group6 = ["设备铭牌及标识", "铭牌和包装", "铭牌", "包装"]
                for kw in keywords_group6:
                    if kw in first_line:
                        return True, [f"我方设备的{first_line}满足上述要求。"], f"标题规则-铭牌类({kw})"

        # ========== 优先级2：基于内容关键词匹配 ==========
        content_rules = [
            ("对产品施工工艺有特殊要求时，投标人应在投标文件中注明", "我方对产品施工工艺无特殊要求。", "内容规则-施工工艺"),
            ("对产品维护有特殊要求时，投标人应在投标文件中注明", "我方对产品维护无特殊要求。", "内容规则-产品维护"),
            ("投标人如有更新的技术，在完成对本招标文件提出的技术条款应答后，可对本招标文件的技术条款提出建议，供招标人参考", "我方暂无更新的技术。", "内容规则-更新技术"),
            ("凡声明提供免费服务的投标人，将被理解为其免费完成招标人的全部服务要求", "凡声明提供免费服务的，将被理解为我方免费完成招标人的全部服务要求。", "内容规则-免费服务"),
            ("除非投标人在投标文件中注明，备品备件的各类参数与原件相同", "我方备品备件的各类参数与原件相同。", "内容规则-备品备件"),
            ("投标人应提供其产品寿命的说明，并提供产品在实际运行中的可靠", "我方将提供其产品寿命的说明，并提供产品在实际运行中的可靠和无维修时间。我方为本项目所供的产品，其寿命为10年，产品在实际运行中的可靠和无维修时间大于20000小时，产品设计寿命将满足技术规格书要求。", "内容规则-产品寿命"),
            ("卖方若采用其它被承认的但没有列在上面的相关国内、国际标准，应明确提出并提供相应标准，经买方批准后方可采用", "我方未采用其他标准。", "内容规则-其他标准"),
            ("所有设备必须出具国家级（CMA、CMC）或国际权威检测机构检查合格报告", "我方所有设备具备国家级（CMA、CMC）或国际权威检测机构检查合格报告。", "内容规则-检测报告"),
        ]
        
        for keyword, response, rule_name in content_rules:
            if keyword in full_text:
                return True, [response], rule_name
        
        return False, None, None

    def process_clauses(self, clause_content_lines):
        """处理条款内容 - 完全按照V2.0原版逻辑"""
        clause_content_lines = [line for line in clause_content_lines if line.strip() != ""]
        
        if not clause_content_lines:
            return []
        
        processed_lines = clause_content_lines.copy()
        first_line = processed_lines[0].strip()

        # 特殊处理规则 - 基于标题关键词匹配
        # 标题判定条件：≤15字符 且 句尾无标点符号
        is_title = len(first_line) <= 15 and not re.search(r'[。！？；：，.!?;:,]$', first_line)

        if is_title:
            # 规则1：项目概况类
            keywords_group1 = ["项目概况", "工程概况", "工程范围", "线路概况", "正线主要技术标准", "联络线技术标准",
                              "招标范围", "适用范围", "使用环境条件", "设计标准", "环境条件", "线路主要技术标准", "设计寿命"]
            for kw in keywords_group1:
                if kw in first_line:
                    return [f"我已了解本项目{first_line}，我方所供产品完全满足本项目{first_line}的要求。"]

            # 规则2：环境地质类
            keywords_group2 = ["使用环境", "地貌特征", "工程地质", "水文地质", "地震动参数", "地震", "海拔", "气象特征", "气象条件"]
            for kw in keywords_group2:
                if kw in first_line:
                    return [f"我已了解本项目{first_line}，我方所供产品完全适用于本项目{first_line}的情况。"]

            # 规则3：设备清单类
            keywords_group3 = ["设备清单", "招标数量", "供货数量及规格", "需求数量", "需求一览表", "备品备件", "专用工具"]
            for kw in keywords_group3:
                if kw in first_line:
                    return [f"我方按照{first_line}要求提供设备，详见商务报价部分。"]

            # 规则4：技术服务
            if "技术服务" in first_line:
                return [f"我方将按照招标文件要求提供{first_line}。"]

            # 规则5：试验类
            keywords_group5 = ["型式试验", "出厂试验", "现场试验"]
            for kw in keywords_group5:
                if kw in first_line:
                    return [f"我方所供设备将按照招标文件的{first_line}的要求来进行试验。"]

            # 规则6：铭牌包装类（仅当条款只有一行时使用简短应答）
            if len(processed_lines) == 1:
                keywords_group6 = ["设备铭牌及标识", "铭牌和包装", "铭牌", "包装"]
                for kw in keywords_group6:
                    if kw in first_line:
                        return [f"我方设备的{first_line}满足上述要求。"]

        # 短条款无标点符号的特殊处理
        line1_name = ['机械外观检查','绝缘试验','功能试验','环境试验','振动试验','电磁兼容试验','连续通电试验']
        if (len(processed_lines) == 1 and              
            (first_line in line1_name or
            first_line.rstrip('；。') in line1_name)):
            # 去掉末尾的分号或句号用于生成回复
            clean_first_line = first_line.rstrip('；。')
            return [f"我方所供设备将按照招标文件的{clean_first_line}的要求来进行{clean_first_line}。"]
        
        # 清理简短首段 - 这是用户提到的关键逻辑！
        if (len(processed_lines) > 1 and 
            len(first_line) <= 30 and 
            not re.search(r'[。！？；：，.!?;]$', first_line)):
            processed_lines[0] = ""
        
        # 标准引用处理 - 完全按照原版
        standard_pattern = re.compile(r'^(GB/T|DL/T|DL/Z|NB/T|TB|TB/T|QC|Q/CR|IEC|GB|TJ/GD)')
        book_title_pattern = re.compile(r'《.+?》')
        end_with_period = re.compile(r'。$')
        standard_paragraphs = [line for line in processed_lines 
                             if standard_pattern.search(line) or book_title_pattern.search(line)]
        
        final_lines = []
        if len(standard_paragraphs) >= 6:
            for line in processed_lines:
                if not (standard_pattern.search(line) or book_title_pattern.search(line)):
                    line = line.replace("下列","上述")
                    
                    line_stripped = line.strip()
                    if line_stripped and not end_with_period.search(line_stripped):
                        if re.search(r'[，,、；;：:！!？?]$', line_stripped):
                            line = re.sub(r'[，,、；;：:！!？?]$', '。', line_stripped)
                        else:
                            line = line_stripped + '。'
                    
                    final_lines.append(line)
        else:
            final_lines = processed_lines
        
        # 处理表格、图片、注释等 - 完全按照原版
        table_figure_pattern = re.compile(r'^表\s*\d+.*|^图\s*\d+.*|^示意图.*')
        end_punctuation = re.compile(r'[。！？；，、.!?;]$')
        note_pattern = re.compile(r'^注\d*[:：]?\s*')
        
        filtered_lines = []
        for line in final_lines:
            line_stripped = line.strip()
            
            if not line_stripped:
                continue
            
            if (table_figure_pattern.match(line_stripped) and 
                not end_punctuation.search(line_stripped) and 
                len(line_stripped) <= 15):
                continue
            
            if note_pattern.match(line_stripped):
                modified_line = note_pattern.sub('', line_stripped)
                if modified_line.strip():
                    filtered_lines.append(modified_line)
                continue
            
            if "示意图" in line_stripped and not end_punctuation.search(line_stripped):
                continue
            
            if "示意图如下" in line_stripped:
                modified_line = line_stripped.replace("示意图如下：", "").replace("示意图如下", "")
                if modified_line.strip() and not re.match(r'^[，,、。；;：:！!？?]+$', modified_line.strip()):
                    filtered_lines.append(modified_line)
                continue
            
            filtered_lines.append(line_stripped)
        
        # 文本替换规则 - 完全按照原版
        override_map = {
            "对产品施工工艺有特殊要求时，投标人应在投标文件中注明。": "我方对产品施工工艺无特殊要求。",
            "对产品维护有特殊要求时，投标人应在投标文件中注明。": "我方对产品维护无特殊要求。",
            "如有偏差，必须提供详细的技术规格偏差表并说明原因。":"我方对技术规格书逐条应答，设备无偏差。",
            "设备应能在使用环境条件下连续运行，如果产品不能满足这些要求，投标人应申报偏差值。":"我方设备能在使用环境条件下连续运行。",
            "投标人如有更新的技术，在完成对本招标文件提出的技术条款应答后，可对本招标文件的技术条款提出建议，供招标人参考。": "我方暂无更新的技术。",
            "凡声明提供免费服务的投标人，将被理解为其免费完成招标人的全部服务要求。":"凡声明提供免费服务的，将被理解为我方免费完成招标人的全部服务要求。",
            "除非投标人在投标文件中注明，备品备件的各类参数与原件相同。":"我方备品备件的各类参数与原件相同。",
            "投标人应提供其产品寿命的说明，并提供产品在实际运行中的可靠和无维修时间。产品设计寿命应满足技术规格书要求。":"我方将提供其产品寿命的说明，并提供产品在实际运行中的可靠和无维修时间。我方为本项目所供的产品，其寿命为10年，产品在实际运行中的可靠和无维修时间大于20000小时，产品设计寿命将满足技术规格书要求。",
            "投标人应提供其产品寿命的说明，并提供产品在实际运行中的可靠和无维修时间。产品设计寿命应满足技术条件要求，无具体要求时，不小于一个大修周期。":"我方将提供其产品寿命的说明，并提供产品在实际运行中的可靠和无维修时间。我方为本项目所供的产品，其寿命为10年，产品在实际运行中的可靠和无维修时间大于20000小时，产品设计寿命将满足技术条件要求，无具体要求时，不小于一个大修周期。",
            "卖方若采用其它被承认的但没有列在上面的相关国内、国际标准，应明确提出并提供相应标准，经买方批准后方可采用。": "我方未采用其他标准。",
            "出具国家级（CMA、CMC）或国际权威检测机构检查合格报告":"具备国家级（CMA、CMC）或国际权威检测机构检查合格报告",
            "出具国家级CMA或CMC或国际权威检测机构的检测报告":"具备国家级CMA或CMC或国际权威检测机构的检测报告"
        }
        
        fixed_sentences = [
            "当采用其它标准时，投标人应在投标书中明确提出各项设备所遵循的标准名称及标准内容并说明与上述标准的差异（需附证明材料）。",
            "当采用其它标准时，投标人应在投标文件中明确提出各项设备所遵循的标准名称及标准内容并说明与上述标准的差异（需附证明材料）。",
            "当采用其它标准时，投标人应在投标文件中明确提出各项设备所遵循的标准名称及内容并说明与上述标准的差异（需附证明材料）。",
            "如有偏差（优于技术规格书要求），必须提供详细的技术规格偏差表。",
            "当采用其它标准时，投标厂商在投标书中明确提出各项设备所遵循的标准名称及标准内容。",
            "或由我方建议的其它等效标准，并提供中文版本，供招标人在招标时确认。",
            "或由投标人建议的其他等效标准，并提供中文版本，由双方在合同文本或设计联络时共同确认。",
            "如本技术条件与上述各标准之间有矛盾，则应满足较高标准的要求。"
        ]
        
        replacement_rules = [
            (r'设备投标商|设备供应商|供应商须|投标[人方商者]|卖方|所有设备厂家', '我方'),
            (r'中标商|中标方|中标人|中标厂商|中标厂家', '若我方中标，我方'),
            (r'应符合但不限于如下标准', '符合如上标准'),
            (r'[但且]不限于', ''),
            (r'以下标准', '以上标准'),
            (r'下述|下列', '上述'),
            (r'应仔细阅读', '已仔细阅读'),
            (r'至少附有', '附有'),
            (r'应该|应当', ''),
            (r'应对', '对'),
            (r'不应有', '无'),
            (r'不应', '不会'),
            (r'应由', '由'),
            (r'应使', '使'),
            (r'应是', '是'),
            (r'应有', '有'),
            (r'宜采用', '采用'),
            (r'(?<![相适对反供响])应(?![答用急变力标])', ''),
            (r'(?<![必所])需(?![要求])', ''),
            (r'(?<!必)须', ''),
            (r'必须', ''),
            (r'需\s*要', '需要'),
            (r'确保\s*其', '保证'),
            (r'应(\w{0,4})要求', r'满足\1要求'),
            (r'^※|※$', '')                           # 只删除行首或行尾的※
        ]
        
        results = []
        for line in filtered_lines:
            for pattern, replacement in override_map.items():
                if pattern in line:
                    line = line.replace(pattern, replacement)
            
            for sentence in fixed_sentences:
                line = line.replace(sentence, "")
            
            for pattern, repl in replacement_rules:
                line = re.sub(pattern, repl, line)
            
            if line.strip():
                results.append(line.strip())
        
        return results

    def add_simple_text(self, paragraph, text):
        """添加简单文本（完全按照V2.0原版逻辑）"""
        # 按字符类型处理文本
        current_run = None
        current_type = None
        
        for ch in text:
            ch_type = 'en' if ord(ch) < 128 else 'cn'
            if ch_type != current_type:
                current_run = paragraph.add_run()
                current_run.bold = True
                
                if ch_type == 'en':
                    current_run.font.name = 'Times New Roman'
                    current_run.font.size = Pt(10.5)
                    try:
                        if current_run._element.rPr is not None:
                            current_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
                    except:
                        pass
                else:
                    current_run.font.name = 'SimSun'
                    current_run.font.size = Pt(10.5)
                    try:
                        if current_run._element.rPr is not None:
                            current_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
                    except:
                        pass
                current_type = ch_type
            
            if current_run is not None:
                current_run.text += ch

    def copy_template_elements(self, target_doc, insert_after_element, response_data):
        """
        从模板文档复制元素（包括段落、表格、图片）到目标文档

        参数：
            target_doc: 目标文档
            insert_after_element: 插入位置（在此元素之后插入）
            response_data: 应答数据字典，包含 template_file 和 element_indices
        返回：
            最后插入的段落元素（用于后续插入）
        """
        import copy
        from lxml import etree

        template_file = response_data.get('template_file')
        element_indices = response_data.get('element_indices', [])

        if not template_file or not element_indices:
            return insert_after_element

        try:
            # 打开模板文档
            template_doc = Document(template_file)
            template_body = list(template_doc._body._body)

            # 先复制所有图片资源，建立关系ID映射
            rel_id_map = self.copy_image_resources_with_mapping(template_doc, target_doc)
            if rel_id_map:
                self.log.emit(f"    已复制 {len(rel_id_map)} 个图片资源，映射: {rel_id_map}")

            last_inserted_element = insert_after_element._p if hasattr(insert_after_element, '_p') else None
            last_para = insert_after_element

            for elem_type, elem_idx in element_indices:
                if elem_idx >= len(template_body):
                    continue

                source_element = template_body[elem_idx]

                # 深度复制元素
                copied_element = copy.deepcopy(source_element)

                # 更新复制元素中的图片关系ID引用
                if rel_id_map:
                    self.update_image_rel_ids(copied_element, rel_id_map)
                    # 检查元素中是否有图片引用
                    elem_xml = copied_element.xml if hasattr(copied_element, 'xml') else str(copied_element)
                    if 'r:embed' in elem_xml or 'drawing' in elem_xml:
                        self.log.emit(f"    元素包含图片引用")

                # 在目标位置之后插入复制的元素
                if last_inserted_element is not None:
                    last_inserted_element.addnext(copied_element)
                else:
                    # 直接添加到文档body
                    target_doc._body._body.append(copied_element)

                # 更新最后插入的元素
                last_inserted_element = copied_element
                tag = copied_element.tag.split('}')[-1] if '}' in copied_element.tag else copied_element.tag
                if tag == 'p':
                    # 为段落创建包装对象
                    last_para = Paragraph(copied_element, target_doc._body)

            # 如果最后插入的不是段落（比如是表格），在其后插入一个空段落作为锚点
            if last_inserted_element is not None:
                tag = last_inserted_element.tag.split('}')[-1] if '}' in last_inserted_element.tag else last_inserted_element.tag
                if tag != 'p':
                    new_p = OxmlElement('w:p')
                    last_inserted_element.addnext(new_p)
                    last_para = Paragraph(new_p, target_doc._body)

            return last_para

        except Exception as e:
            self.log.emit(f"复制模板元素失败: {e}")
            import traceback
            traceback.print_exc()
            return insert_after_element

    def copy_image_resources_with_mapping(self, source_doc, target_doc):
        """
        复制源文档中的图片资源到目标文档，并返回关系ID映射

        返回：
            dict: {旧关系ID: 新关系ID}
        """
        from docx.opc.packuri import PackURI
        from docx.parts.image import ImagePart
        import hashlib

        rel_id_map = {}

        try:
            source_part = source_doc.part
            target_part = target_doc.part
            target_package = target_doc.part.package

            # 用于跟踪已添加的图片（避免重复）
            added_images = {}  # {图片hash: new_rel_id}

            # 复制所有图片关系
            for old_rel_id, rel in source_part.rels.items():
                if "image" in rel.reltype:
                    try:
                        # 获取图片数据
                        image_part = rel.target_part
                        image_blob = image_part.blob
                        content_type = image_part.content_type

                        # 计算图片hash，避免重复添加
                        image_hash = hashlib.md5(image_blob).hexdigest()

                        if image_hash in added_images:
                            # 已经添加过这个图片，直接使用已有的关系ID
                            rel_id_map[old_rel_id] = added_images[image_hash]
                            continue

                        # 获取图片扩展名
                        partname = str(image_part.partname)
                        ext = partname.split('.')[-1].lower() if '.' in partname else 'png'

                        # 生成新的图片路径
                        existing_images = [p.partname for p in target_package.parts if '/media/image' in str(p.partname)]
                        next_num = len(existing_images) + 1
                        new_partname = f'/word/media/image{next_num}.{ext}'

                        # 创建新的图片部件
                        new_image_part = ImagePart.load(PackURI(new_partname), content_type, image_blob, target_package)

                        # 添加到目标包
                        target_package.parts.append(new_image_part)

                        # 创建关系
                        new_rel_id = target_part.relate_to(new_image_part, rel.reltype)
                        rel_id_map[old_rel_id] = new_rel_id
                        added_images[image_hash] = new_rel_id

                    except Exception as img_error:
                        self.log.emit(f"复制图片资源失败: {img_error}")
                        import traceback
                        traceback.print_exc()

        except Exception as e:
            self.log.emit(f"复制图片资源整体失败: {e}")
            import traceback
            traceback.print_exc()

        return rel_id_map

    def update_image_rel_ids(self, element, rel_id_map):
        """
        更新元素中的图片关系ID引用

        参数：
            element: XML元素
            rel_id_map: {旧关系ID: 新关系ID} 映射
        """
        if not rel_id_map:
            return

        # 定义命名空间
        nsmap = {
            'a': 'http://schemas.openxmlformats.org/drawingml/2006/main',
            'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships',
            'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing',
        }

        # 查找所有包含 r:embed 属性的元素（图片引用）
        for attr_name in ['{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed',
                          '{http://schemas.openxmlformats.org/officeDocument/2006/relationships}link']:
            for elem in element.iter():
                old_id = elem.get(attr_name)
                if old_id and old_id in rel_id_map:
                    elem.set(attr_name, rel_id_map[old_id])

    def insert_paragraph_after(self, paragraph, text=None, style=None):
        """在段落后插入新段落"""
        new_p = OxmlElement('w:p')
        paragraph._p.addnext(new_p)
        new_para = Paragraph(new_p, paragraph._parent)
        if text:
            new_para.add_run(text)
        if style is not None:
            new_para.style = style
        return new_para

    def remove_only_pure_empty_paragraphs(self, doc):
        """删除纯空段落"""
        self.file_logger.info('[remove_only_pure_empty_paragraphs] 开始')
        deleted_count = 0

        for para in reversed(doc.paragraphs):
            if self.is_pure_empty_paragraph(para):
                para._element.getparent().remove(para._element)
                deleted_count += 1

        self.log.emit(f"已删除 {deleted_count} 个空段落")
        self.file_logger.info(f'[remove_only_pure_empty_paragraphs] 删除 {deleted_count} 个空段落')
    
    def is_pure_empty_paragraph(self, paragraph):
        """判断是否为纯空段落"""
        has_no_text = not paragraph.text.strip()

        xml = paragraph._p.xml
        # 同时检测 w:drawing (新格式) 和 w:pict (旧格式VML)
        has_no_special_elements = all(
            tag not in xml
            for tag in ['<w:drawing', '<w:pict', '<w:tbl', '<w:br', '<w:sectPr']
        )

        return has_no_text and has_no_special_elements